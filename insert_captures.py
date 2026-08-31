"""
Insert captures d'ecran dans Memoir_FINAL_V22.docx → Memoir_FINAL_V23.docx

Usage:
    python insert_captures.py

Prerequis:
    pip install python-docx Pillow
"""
import os
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'C:\Users\loulou\Downloads\Memoir_FINAL_V22.docx'
DST = r'C:\Users\loulou\Downloads\Memoir_FINAL_V23.docx'
IMG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots-memoire')

PLACEHOLDER_MAP = {
    'Visual Studio Code': 'iv3_vscode.png',
    'XAMPP et phpMyAdmin': 'iv3_phpmyadmin.png',
    'Postman': 'iv3_postman.png',
    'Android Studio': 'iv3_android_studio.png',
    'GitHub': 'iv3_github.png',
    'Claude Code CLI': None,
}

def insert_image_at_paragraph(doc, para_idx, img_path, width_cm=15):
    p = doc.paragraphs[para_idx]
    p.clear()
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def main():
    if not os.path.exists(SRC):
        print(f'ERREUR: {SRC} introuvable')
        return

    doc = Document(SRC)
    total = len(doc.paragraphs)
    replaced = 0
    missing = []

    for i in range(total):
        text = doc.paragraphs[i].text.strip()

        for label, filename in PLACEHOLDER_MAP.items():
            placeholder = f"[Insérer ici la capture d'écran : {label}]"
            if placeholder in text or (label in text and 'Insérer' in text):
                if filename is None:
                    print(f'  [{i}] {label} — pas de capture definie, placeholder conserve')
                    missing.append(label)
                    continue

                img_path = os.path.join(IMG_DIR, filename)
                if not os.path.exists(img_path):
                    print(f'  [{i}] {label} — MANQUANTE ({filename})')
                    missing.append(f'{label} ({filename})')
                    continue

                insert_image_at_paragraph(doc, i, img_path)
                replaced += 1
                print(f'  [{i}] {label} — insere ({filename})')
                break

    doc.save(DST)
    print(f'\n{"="*50}')
    print(f'Captures inserees : {replaced}')
    print(f'Manquantes        : {len(missing)}')
    for m in missing:
        print(f'  - {m}')
    print(f'\nSauvegarde : {DST}')

if __name__ == '__main__':
    main()
