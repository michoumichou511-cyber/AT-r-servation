# -*- coding: utf-8 -*-
"""Génère docs/Chapitre2_Analyse_Conception.docx (texte structuré, prêt à l'impression)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DOCS = Path(__file__).resolve().parent
OUT = DOCS / "Chapitre2_Analyse_Conception.docx"

AT_GREEN = RGBColor(0x00, 0xA6, 0x50)
AT_BLUE = RGBColor(0x00, 0x3D, 0xA5)


def add_title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = AT_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.runs[0].italic = True
        s.runs[0].font.size = Pt(11)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(
        doc,
        "Chapitre 2 — Analyse et conception",
        "Projet de fin de formation — AT Réservations (Algérie Télécom)",
    )

    intro = (
        "Ce chapitre présente l'analyse fonctionnelle et la conception du système d'information "
        "AT Réservations, plateforme web de gestion des missions et déplacements professionnels. "
        "Stack : React 18 (interface), Laravel 12 (API) et MySQL (persistance). "
        "Quatre rôles : administrateur, validateur, demandeur et utilisateur.\n\n"
        "Les diagrammes UML officiels sont fournis en source PlantUML dans le fichier "
        "diagrammes.puml (même dossier docs/), pour export vectoriel ou PNG sans perte de qualité."
    )
    doc.add_paragraph(intro)

    doc.add_heading("2.1 Diagramme de cas d'utilisation", level=1)
    doc.add_paragraph(
        "Objectif : représenter les interactions entre acteurs et système.\n"
        "Acteurs : Administrateur, Validateur, Demandeur, Utilisateur.\n"
        "Cas : s'authentifier ; créer une demande de mission ; soumettre la demande ; "
        "valider ou rejeter une demande ; consulter mes missions ; gérer les utilisateurs ; "
        "consulter le tableau de bord ; exporter un rapport ; envoyer un message interne ; "
        "consulter l'organigramme.\n"
        "Relation include : soumettre complète la création lorsque le demandeur valide son dossier."
    )
    p = doc.add_paragraph()
    r = p.add_run("Source PlantUML : bloc diagramme_cas_utilisation")
    r.italic = True
    r.font.color.rgb = AT_GREEN

    doc.add_heading("2.2 Diagramme de classes", level=1)
    doc.add_paragraph(
        "Classes : Utilisateur (id, nom, prenom, email, role, direction) ; "
        "Mission (id, titre, destination, dates, statut, type_mission, description, created_by) ; "
        "Validation (mission_id, validateur_id, statut, commentaire, date_validation) ; "
        "Document (mission_id, nom_fichier, chemin, type) ; "
        "Message (expediteur_id, destinataire_id, contenu, lu, created_at) ; "
        "Notification (user_id, message, type, lu, created_at).\n"
        "Relations : un utilisateur crée plusieurs missions ; une mission a plusieurs validations "
        "et documents ; messages et notifications liés aux utilisateurs avec cardinalités 1-*."
    )
    p = doc.add_paragraph()
    r = p.add_run("Source PlantUML : bloc diagramme_classes")
    r.italic = True
    r.font.color.rgb = AT_GREEN

    doc.add_heading("2.3 Diagramme de séquence — Création et validation", level=1)
    doc.add_paragraph(
        "Participants : Demandeur, Frontend React, API Laravel, MySQL, Validateur.\n"
        "Séquence : authentification ; remplissage et soumission du formulaire ; "
        "enregistrement mission ; notification au validateur ; décision (valider) ; "
        "notification au demandeur ; mission approuvée."
    )
    p = doc.add_paragraph()
    r = p.add_run("Source PlantUML : bloc diagramme_sequence_mission")
    r.italic = True
    r.font.color.rgb = AT_GREEN

    doc.add_heading("2.4 Diagramme d'activités — Circuit hiérarchique", level=1)
    doc.add_paragraph(
        "Du brouillon à l'archivage : soumission, file d'attente validateur, "
        "branches approuvé / rejeté / retour pour modification, notifications, "
        "resoumission éventuelle, archivage en fin de cycle."
    )
    p = doc.add_paragraph()
    r = p.add_run("Source PlantUML : bloc diagramme_activite_validation")
    r.italic = True
    r.font.color.rgb = AT_GREEN

    doc.add_heading("2.5 Diagramme de déploiement", level=1)
    doc.add_paragraph(
        "Client : navigateur + bundle React (build Vite). "
        "Serveur : Apache/Nginx + PHP-FPM, application Laravel 12. "
        "Données : MySQL. Flux HTTPS entre SPA et API REST."
    )
    p = doc.add_paragraph()
    r = p.add_run("Source PlantUML : bloc diagramme_deploiement")
    r.italic = True
    r.font.color.rgb = AT_GREEN

    doc.add_heading("Annexe — Export des figures", level=1)
    doc.add_paragraph(
        "Pour insérer les schémas dans ce document Word : générer les PNG ou SVG "
        "avec PlantUML (voir commande dans GENERER_DIAGRAMMES.sh), puis Insertion > Images dans Word. "
        "Les couleurs charte AT utilisées dans les diagrammes sont #00A650 (vert) et #003DA5 (bleu)."
    )

    doc.save(OUT)
    print(f"Écrit : {OUT}")


if __name__ == "__main__":
    main()
