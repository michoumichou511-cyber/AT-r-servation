# -*- coding: utf-8 -*-
# Assemblage : Chapitre II (rédigé) + Chapitre III (intégré avec figures)
import copy
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r"C:\Users\loulou\ProjetFinFormation\Memoir_BASE.docx"
CAP3 = r"C:\Users\loulou\ProjetFinFormation\captures_memoire\chap3"
DIAG = r"C:\Users\loulou\ProjetFinFormation\diagrammes"
LOGO = r"C:\Users\loulou\ProjetFinFormation\logos\logo_algerie_telecom.png"

d = Document(BASE)

def para(text="", style=None, align=None):
    p = d.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    return p

def image(path, width_cm=15.5):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    return p

def caption(text):
    p = para(text, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        r.font.italic = True
        r.font.bold = True
        r.font.size = Pt(10)
    return p

def bullets(items):
    for it in items:
        bp = d.add_paragraph(it, style="List Paragraph")
        bp.paragraph_format.left_indent = Cm(1)

def uc_table(rows):
    t = d.add_table(rows=len(rows), cols=2)
    t.style = d.tables[1].style
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for r in t.rows[i].cells[0].paragraphs[0].runs:
            r.font.bold = True
    return t

# ================= CHAPITRE II =================
para("CHAPITRE II", style="Heading 1")
para("Analyse et Spécification des Besoins", style="Heading 2")

para("Introduction", style="Heading 3")
para("Après avoir présenté l'organisme d'accueil et le contexte général du projet dans le chapitre précédent, ce chapitre est consacré à l'analyse et à la spécification des besoins de l'application « AT Réservations ». Nous y identifions les différents acteurs du système, leurs besoins fonctionnels, ainsi que les besoins non fonctionnels auxquels l'application doit répondre. Cette analyse est ensuite formalisée à l'aide du diagramme de cas d'utilisation UML, complété par les descriptions textuelles des cas d'utilisation les plus importants.")

para("Identification des acteurs", style="Heading 3")
para("Un acteur représente une entité externe (utilisateur ou système) qui interagit directement avec l'application. Quatre acteurs principaux ont été identifiés :")
bullets([
    "Le Demandeur : tout agent d'Algérie Télécom habilité à créer, suivre et gérer ses demandes de mission ;",
    "Le Directeur : responsable hiérarchique chargé de valider ou de refuser les demandes de mission de sa direction ;",
    "L'Agent DML : agent de la Direction des Moyens Logistiques chargé de la prise en charge logistique des missions validées (hôtels, véhicules, bons de transport) ;",
    "L'Administrateur : agent de la DSI chargé de la gestion des comptes, des référentiels (hôtels, véhicules, budgets) et de la supervision globale (espace web uniquement).",
])

para("Les besoins fonctionnels", style="Heading 3")
para("Les besoins fonctionnels expriment les services que le système doit rendre à ses utilisateurs. Ils sont organisés par acteur :")
para("Pour le Demandeur :", style="Heading 4")
bullets([
    "S'authentifier de manière sécurisée ;",
    "Gérer ses missions : créer une demande, la modifier tant qu'elle n'est pas validée, la supprimer ;",
    "Suivre l'état d'avancement de ses missions en temps réel ;",
    "Recevoir des notifications sur l'évolution de ses demandes de mission.",
])
para("Pour le Directeur :", style="Heading 4")
bullets([
    "S'authentifier de manière sécurisée ;",
    "Recevoir les demandes de mission par une notification ;",
    "Valider les demandes de mission : les approuver ou les refuser avec motif ;",
    "Consulter l'historique des missions et exporter une mission en PDF ;",
    "Envoyer des messages internes aux autres utilisateurs.",
])
para("Pour l'Agent DML :", style="Heading 4")
bullets([
    "S'authentifier de manière sécurisée ;",
    "Recevoir des notifications sur les missions validées ;",
    "Consulter les missions validées en attente de traitement ;",
    "Consulter la liste des hôtels (conventionnés ou non conventionnés) ;",
    "Affecter un hôtel à une mission ;",
    "Affecter des véhicules aux missions, après vérification de leur disponibilité.",
])

para("Les besoins non fonctionnels", style="Heading 3")
para("Au-delà des fonctionnalités, l'application doit satisfaire des exigences de qualité :")
bullets([
    "Sécurité : authentification Active Directory/LDAP avec repli base de données, contrôle d'accès par rôles (RBAC), mots de passe hachés (bcrypt), traçabilité par journal d'audit ;",
    "Performance : temps de réponse inférieur à 2 secondes pour les opérations courantes ;",
    "Ergonomie : interface intuitive, responsive et conforme à la charte graphique d'Algérie Télécom ;",
    "Disponibilité : accessibilité permanente sur l'intranet de l'entreprise ;",
    "Maintenabilité : architecture en couches (API REST Laravel / client React / mobile Flutter) facilitant les évolutions ;",
    "Portabilité : application web multi-navigateurs et application mobile multi-plateformes (Android/iOS).",
])

para("Diagramme de cas d'utilisation", style="Heading 3")
para("Le diagramme de cas d'utilisation ci-dessous synthétise les interactions entre les trois acteurs métier (Demandeur, Agent DML, Directeur) et le système. Chaque cas d'utilisation inclut l'authentification préalable (relation « include ») ; les relations « extend » représentent les variantes optionnelles d'un cas de base.")
image(DIAG + r"\cas_utilisation.png", 16.0)
caption("Figure N°3 : Diagramme de cas d'utilisation global de l'application AT Réservations")

para("Description textuelle des cas d'utilisation", style="Heading 3")
para("Les fiches suivantes détaillent les trois cas d'utilisation majeurs du système.")

para("Cas d'utilisation « Créer une mission »", style="Heading 4")
uc_table([
    ("Nom", "Créer une mission"),
    ("Acteur principal", "Demandeur"),
    ("Précondition", "Le demandeur est authentifié."),
    ("Scénario nominal", "1. Le demandeur accède au formulaire de nouvelle mission.\n2. Il renseigne le titre, la destination, les dates, le motif et les besoins (hébergement, transport).\n3. Il joint les pièces justificatives éventuelles.\n4. Il soumet la demande.\n5. Le système enregistre la demande, notifie le directeur et affiche une confirmation."),
    ("Scénarios alternatifs", "2a. Champs invalides : le système signale les erreurs et le demandeur corrige.\n4a. Le demandeur enregistre la demande en brouillon pour la compléter plus tard."),
    ("Postcondition", "La demande est créée avec le statut « Soumise » et entre dans le circuit de validation."),
])
para("Tableau N°2 : Description textuelle du cas « Créer une mission »", align=WD_ALIGN_PARAGRAPH.CENTER)

para("Cas d'utilisation « Valider une demande de mission »", style="Heading 4")
uc_table([
    ("Nom", "Valider une demande de mission"),
    ("Acteur principal", "Directeur"),
    ("Précondition", "Le directeur est authentifié ; une demande est en attente de validation dans son périmètre."),
    ("Scénario nominal", "1. Le directeur consulte la liste des demandes en attente.\n2. Il ouvre le détail d'une demande.\n3. Il approuve la demande.\n4. Le système change le statut, notifie le demandeur et transmet la mission à l'agent DML."),
    ("Scénarios alternatifs", "3a. Le directeur refuse la demande : il saisit un motif obligatoire et le système notifie le demandeur.\n3b. Le directeur demande une modification : la demande retourne au demandeur."),
    ("Postcondition", "La demande est approuvée (ou refusée) et tracée dans le journal d'audit."),
])
para("Tableau N°3 : Description textuelle du cas « Valider une demande de mission »", align=WD_ALIGN_PARAGRAPH.CENTER)

para("Cas d'utilisation « Traiter la logistique d'une mission »", style="Heading 4")
uc_table([
    ("Nom", "Traiter la logistique d'une mission"),
    ("Acteur principal", "Agent DML"),
    ("Précondition", "L'agent DML est authentifié ; une mission validée est en attente de traitement."),
    ("Scénario nominal", "1. L'agent DML consulte les missions validées à traiter.\n2. Il affecte un hôtel conventionné selon la destination.\n3. Il affecte un moyen de transport (véhicule de service disponible, billet…).\n4. Il marque la logistique comme « OK ».\n5. Le système notifie le demandeur et clôture le volet logistique."),
    ("Scénarios alternatifs", "2a. Aucun hôtel conventionné disponible : l'agent renseigne un hôtel hors convention.\n3a. Aucun véhicule disponible : l'agent émet un bon de transport."),
    ("Postcondition", "La mission passe au statut « Logistique OK » ; toutes les affectations sont enregistrées."),
])
para("Tableau N°4 : Description textuelle du cas « Traiter la logistique »", align=WD_ALIGN_PARAGRAPH.CENTER)

para("Conclusion", style="Heading 3")
para("Ce chapitre a permis d'identifier les acteurs du système, de recenser les besoins fonctionnels et non fonctionnels, et de les formaliser au moyen du diagramme de cas d'utilisation et des descriptions textuelles. Cette spécification constitue le socle de la phase de conception présentée dans le chapitre suivant.")

# ================= CHAPITRE III (intégration) =================
src = Document(r"C:\Users\loulou\ProjetFinFormation\ChapitreIII_src.docx")
img_map = {
    "logo d'Algérie Télécom": (LOGO, 6.0),
    "trame générale": (CAP3 + r"\fig02_trame.png", 15.5),
    "Page de connexion": (CAP3 + r"\fig03_login.png", 15.5),
    "Espace Administrateur]": (CAP3 + r"\fig04_dash_admin.png", 15.5),
    "Espace Directeur]": (CAP3 + r"\fig05_dash_directeur.png", 15.5),
    "Espace Demandeur]": (CAP3 + r"\fig06_dash_demandeur.png", 15.5),
    "Espace Agent DML]": (CAP3 + r"\fig07_dash_dml.png", 15.5),
    "Gestion des utilisateurs": (CAP3 + r"\fig08_utilisateurs.png", 15.5),
    "Création d'une demande": (CAP3 + r"\fig09_creation_mission.png", 15.5),
    "Circuit de validation": (CAP3 + r"\fig10_validations.png", 15.5),
    "Prise en charge logistique": (CAP3 + r"\fig11_dml_logistique.png", 15.5),
    "conventions hôtelières": (CAP3 + r"\fig12_conventions.png", 15.5),
    "parc véhicules": (CAP3 + r"\fig13_vehicules.png", 15.5),
    "Page Statistiques": (CAP3 + r"\fig14_statistiques.png", 15.5),
    "Messagerie interne": (CAP3 + r"\fig15_messagerie.png", 15.5),
    "Organigramme interactif": (CAP3 + r"\fig16_organigramme.png", 15.5),
    "diagramme de classes": (DIAG + r"\diagramme_classes.png", 15.0),
}

inserted = 0
for el in list(src.element.body):
    tag = el.tag.split("}")[1]
    if tag == "sectPr":
        continue
    if tag == "p":
        text = "".join(n.text or "" for n in el.iter(qn("w:t"))).strip()
        if text.startswith("[Insérer ici"):
            hit = next((v for k, v in img_map.items() if k in text), None)
            if hit:
                image(hit[0], hit[1])
                inserted += 1
            else:
                print("PLACEHOLDER SANS IMAGE:", text[:80])
            continue
    d.element.body.append(copy.deepcopy(el))

print("images inserees dans chap III:", inserted)

# sauts de page + docPr uniques
for p in d.paragraphs:
    if p.text.strip() in ("CHAPITRE II", "CHAPITRE III"):
        p.paragraph_format.page_break_before = True
nid = 100
for docPr in d.element.body.iter(qn("wp:docPr")):
    docPr.set("id", str(nid))
    docPr.set("name", "Image %d" % nid)
    nid += 1

d.save(BASE)
print("assemblage termine, images totales:", nid - 100)
