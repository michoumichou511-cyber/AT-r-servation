# -*- coding: utf-8 -*-
"""Assemblage du memoire MASTER (base FINAL_1) :
couverture, espaces, chapitres II/III/IV, conclusion, bibliographie,
sections par chapitre avec renvois d'en-tetes (rId100-102, parts ajoutees ensuite).
"""
import copy
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE  = r"C:\Users\loulou\ProjetFinFormation\Memoir_MASTER.docx"
CAP3  = r"C:\Users\loulou\ProjetFinFormation\captures_memoire\chap3"
DIAG  = r"C:\Users\loulou\ProjetFinFormation\diagrammes"
SCEAU = r"C:\Users\loulou\ProjetFinFormation\logos\sceau_ministere.png"
AT    = r"C:\Users\loulou\ProjetFinFormation\logos\logo_algerie_telecom.png"

d = Document(BASE)

# ══════════════ 1. COUVERTURE ══════════════
paras = d.paragraphs
p_rep = next(p for p in paras if p.text.strip().startswith("LA REPUBLIQUE"))
i_rep = paras.index(p_rep)
header_ps = paras[i_rep:i_rep + 3]
tbl = d.add_table(rows=1, cols=3)
tbl.autofit = False
for cell, w in zip(tbl.rows[0].cells, (Cm(3.2), Cm(10.0), Cm(3.2))):
    cell.width = w
header_ps[0]._p.addprevious(tbl._tbl)
left, mid, right = tbl.rows[0].cells
for cell in (left, right):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run().add_picture(SCEAU, width=Cm(2.6))
mid._tc.remove(mid.paragraphs[0]._p)
for p in header_ps:
    mid._tc.append(p._p)

# bordure basse du cadre theme
last_theme = next(p for p in d.paragraphs if "AT Réservations »" in p.text)
first_theme = next(p for p in d.paragraphs if p.text.strip().startswith("Conception Et"))
top = first_theme._p.find(qn('w:pPr')).find(qn('w:pBdr')).find(qn('w:top'))
pBdr_last = last_theme._p.find(qn('w:pPr')).find(qn('w:pBdr'))
if pBdr_last.find(qn('w:bottom')) is None:
    bottom = copy.deepcopy(top)
    bottom.tag = qn('w:bottom')
    pBdr_last.append(bottom)

# logo AT a cote d'Organisme d'accueil
p_org = next(p for p in d.paragraphs if p.text.strip().startswith("Organisme d'accueil"))
if not p_org._p.findall('.//' + qn('w:drawing')):
    r = p_org.add_run("  ")
    r.add_picture(AT, height=Cm(1.1))
print("couverture OK")

# ══════════════ 2. NETTOYAGE DES ESPACES ══════════════
def is_empty(p):
    return (not p.text.strip()
            and not p._p.findall('.//' + qn('w:drawing')))

def has_sectpr(p):
    pPr = p._p.find(qn('w:pPr'))
    return pPr is not None and pPr.find(qn('w:sectPr')) is not None

paras = d.paragraphs
i_rem = next(i for i, p in enumerate(paras) if p.text.strip() == "Remerciements")
removed = 0
plist = d.paragraphs
i = i_rem
while i < len(plist):
    if is_empty(plist[i]) and not has_sectpr(plist[i]):
        j = i
        while j < len(plist) and is_empty(plist[j]) and not has_sectpr(plist[j]):
            j += 1
        nxt = plist[j].style.name if j < len(plist) else ""
        keep = 0 if "Heading" in nxt else 1
        for k in range(i + keep, j):
            plist[k]._p.getparent().remove(plist[k]._p)
            removed += 1
        plist = d.paragraphs
        i = i + keep + 1
    else:
        i += 1
print("vides supprimes:", removed)

# ══════════════ HELPERS ══════════════
def para(text="", style=None, align=None):
    p = d.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    return p

def image(path, width_cm=15.5):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    return p

def caption(text):
    p = para(text, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        r.font.italic = True
        r.font.bold = True
        r.font.size = Pt(10)
    return p

def bullets(items):
    for it in items:
        bp = d.add_paragraph(it, style="List Paragraph")
        bp.paragraph_format.left_indent = Cm(1)

def kv_table(rows, bold_col0=True):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = next(tt for tt in d.tables if len(tt.columns) == 2 and "Nom de l'entreprise" in tt.rows[0].cells[0].text).style
    for i, row in enumerate(rows):
        for jj, val in enumerate(row):
            t.rows[i].cells[jj].text = val
        if bold_col0:
            for r in t.rows[i].cells[0].paragraphs[0].runs:
                r.font.bold = True
    return t

BODY_SECT = d.element.body.find(qn('w:sectPr'))

def section_break(header_rid):
    """Paragraphe vide portant un sectPr clone du sectPr final,
    avec en-tete = header_rid et pied = footer3 (rId11, numeros de page)."""
    p = d.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = copy.deepcopy(BODY_SECT)
    for c in list(sp):
        if 'Reference' in c.tag:
            sp.remove(c)
    hr = OxmlElement('w:headerReference')
    hr.set(qn('w:type'), 'default')
    hr.set(qn('r:id'), header_rid)
    fr = OxmlElement('w:footerReference')
    fr.set(qn('w:type'), 'default')
    fr.set(qn('r:id'), 'rId11')
    sp.insert(0, fr)
    sp.insert(0, hr)
    pPr.append(sp)
    return p

# ══════════════ 3. CHAPITRE II ══════════════
para("CHAPITRE II", style="Heading 1")
para("Analyse et Spécification des Besoins", style="Heading 2")
para("Introduction", style="Heading 3")
para("Après avoir présenté l'organisme d'accueil et le contexte général du projet dans le chapitre précédent, ce chapitre est consacré à l'analyse et à la spécification des besoins de l'application « AT Réservations ». Nous y identifions les différents acteurs du système, leurs besoins fonctionnels, ainsi que les besoins non fonctionnels auxquels l'application doit répondre. Cette analyse est ensuite formalisée à l'aide du diagramme de cas d'utilisation UML, complété par les descriptions textuelles des cas d'utilisation les plus importants.")
para("Identification des acteurs", style="Heading 3")
para("Un acteur représente une entité externe (utilisateur ou système) qui interagit directement avec l'application. Quatre acteurs principaux ont été identifiés :")
bullets([
    "Le Demandeur : tout agent d'Algérie Télécom habilité à créer, suivre et gérer ses demandes de mission ;",
    "Le Directeur : responsable hiérarchique chargé de valider ou de refuser les demandes de mission de sa direction ;",
    "L'Agent DML : agent de la Direction des Moyens Logistiques chargé de la prise en charge logistique des missions validées (hôtels, véhicules, bons de transport) ;",
    "L'Administrateur : agent de la DSI chargé de la gestion des comptes, des référentiels (hôtels, véhicules, budgets) et de la supervision globale (rôle système, espace web uniquement).",
])
para("Les besoins fonctionnels", style="Heading 3")
para("Les besoins fonctionnels expriment les services que le système doit rendre à ses utilisateurs. Ils sont organisés par acteur :")
para("Pour le Demandeur :", style="Heading 4")
bullets([
    "S'authentifier de manière sécurisée ;",
    "Gérer ses missions : créer une demande, la modifier tant qu'elle n'est pas validée, la supprimer ;",
    "Suivre l'état d'avancement de ses missions en temps réel ;",
    "Recevoir des notifications sur l'évolution de ses demandes de mission.",
])
para("Pour le Directeur :", style="Heading 4")
bullets([
    "S'authentifier de manière sécurisée ;",
    "Recevoir les demandes de mission par une notification ;",
    "Valider les demandes de mission : les approuver ou les refuser avec motif ;",
    "Consulter l'historique des missions et exporter une mission en PDF ;",
    "Envoyer des messages internes aux autres utilisateurs.",
])
para("Pour l'Agent DML :", style="Heading 4")
bullets([
    "S'authentifier de manière sécurisée ;",
    "Recevoir des notifications sur les missions validées ;",
    "Consulter les missions validées en attente de traitement ;",
    "Consulter la liste des hôtels (conventionnés ou non conventionnés) ;",
    "Affecter un hôtel à une mission ;",
    "Affecter des véhicules aux missions, après vérification de leur disponibilité.",
])
para("Les besoins non fonctionnels", style="Heading 3")
para("Au-delà des fonctionnalités, l'application doit satisfaire des exigences de qualité :")
bullets([
    "Sécurité : authentification Active Directory/LDAP avec repli base de données, contrôle d'accès par rôles (RBAC), mots de passe hachés (bcrypt), traçabilité par journal d'audit ;",
    "Performance : temps de réponse inférieur à 2 secondes pour les opérations courantes ;",
    "Ergonomie : interface intuitive, responsive et conforme à la charte graphique d'Algérie Télécom ;",
    "Disponibilité : accessibilité permanente sur l'intranet de l'entreprise ;",
    "Maintenabilité : architecture en couches (API REST Laravel / client React / mobile Flutter) facilitant les évolutions ;",
    "Portabilité : application web multi-navigateurs et application mobile multi-plateformes (Android/iOS).",
])
para("Diagramme de cas d'utilisation", style="Heading 3")
para("Le diagramme de cas d'utilisation ci-dessous synthétise les interactions entre les trois acteurs métier (Demandeur, Agent DML, Directeur) et le système. Chaque cas d'utilisation inclut l'authentification préalable (relation « include ») ; les relations « extend » représentent les variantes optionnelles d'un cas de base.")
image(DIAG + r"\cas_utilisation.png", 16.0)
caption("Figure N°3 : Diagramme de cas d'utilisation global de l'application AT Réservations")
para("Description textuelle des cas d'utilisation", style="Heading 3")
para("Les fiches suivantes détaillent les trois cas d'utilisation majeurs du système.")
para("Cas d'utilisation « Créer une mission »", style="Heading 4")
kv_table([
    ("Nom", "Créer une mission"),
    ("Acteur principal", "Demandeur"),
    ("Précondition", "Le demandeur est authentifié."),
    ("Scénario nominal", "1. Le demandeur accède au formulaire de nouvelle mission.\n2. Il renseigne le titre, la destination, les dates, le motif et les besoins (hébergement, transport).\n3. Il joint les pièces justificatives éventuelles.\n4. Il soumet la demande.\n5. Le système enregistre la demande, notifie le directeur et affiche une confirmation."),
    ("Scénarios alternatifs", "2a. Champs invalides : le système signale les erreurs et le demandeur corrige.\n4a. Le demandeur enregistre la demande en brouillon pour la compléter plus tard."),
    ("Postcondition", "La demande est créée avec le statut « Soumise » et entre dans le circuit de validation."),
])
caption("Tableau N°2 : Description textuelle du cas « Créer une mission »")
para("Cas d'utilisation « Valider une demande de mission »", style="Heading 4")
kv_table([
    ("Nom", "Valider une demande de mission"),
    ("Acteur principal", "Directeur"),
    ("Précondition", "Le directeur est authentifié ; une demande est en attente de validation dans son périmètre."),
    ("Scénario nominal", "1. Le directeur consulte la liste des demandes en attente.\n2. Il ouvre le détail d'une demande.\n3. Il approuve la demande.\n4. Le système change le statut, notifie le demandeur et transmet la mission à l'agent DML."),
    ("Scénarios alternatifs", "3a. Le directeur refuse la demande : il saisit un motif obligatoire et le système notifie le demandeur.\n3b. Le directeur demande une modification : la demande retourne au demandeur."),
    ("Postcondition", "La demande est approuvée (ou refusée) et tracée dans le journal d'audit."),
])
caption("Tableau N°3 : Description textuelle du cas « Valider une demande de mission »")
para("Cas d'utilisation « Traiter la logistique d'une mission »", style="Heading 4")
kv_table([
    ("Nom", "Traiter la logistique d'une mission"),
    ("Acteur principal", "Agent DML"),
    ("Précondition", "L'agent DML est authentifié ; une mission validée est en attente de traitement."),
    ("Scénario nominal", "1. L'agent DML consulte les missions validées à traiter.\n2. Il affecte un hôtel conventionné selon la destination.\n3. Il affecte un moyen de transport (véhicule de service disponible, billet…).\n4. Il marque la logistique comme « OK ».\n5. Le système notifie le demandeur et clôture le volet logistique."),
    ("Scénarios alternatifs", "2a. Aucun hôtel conventionné disponible : l'agent renseigne un hôtel hors convention.\n3a. Aucun véhicule disponible : l'agent émet un bon de transport."),
    ("Postcondition", "La mission passe au statut « Logistique OK » ; toutes les affectations sont enregistrées."),
])
caption("Tableau N°4 : Description textuelle du cas « Traiter la logistique »")
para("Conclusion", style="Heading 3")
para("Ce chapitre a permis d'identifier les acteurs du système, de recenser les besoins fonctionnels et non fonctionnels, et de les formaliser au moyen du diagramme de cas d'utilisation et des descriptions textuelles. Cette spécification constitue le socle de la phase de conception présentée dans le chapitre suivant.")

section_break('rId100')   # fin du chapitre II -> en-tete Chapitre II
print("chapitre II OK")

# ══════════════ 4. CHAPITRE III (import + renumerotation) ══════════════
src = Document(r"C:\Users\loulou\ProjetFinFormation\ChapitreIII_src.docx")
img_map = {
    "logo d'Algérie Télécom": (AT, 6.0),
    "trame générale": (CAP3 + r"\fig02_trame.png", 15.5),
    "Page de connexion": (CAP3 + r"\fig03_login.png", 15.5),
    "Espace Administrateur]": (CAP3 + r"\fig04_dash_admin.png", 15.5),
    "Espace Directeur]": (CAP3 + r"\fig05_dash_directeur.png", 15.5),
    "Espace Demandeur]": (CAP3 + r"\fig06_dash_demandeur.png", 15.5),
    "Espace Agent DML]": (CAP3 + r"\fig07_dash_dml.png", 15.5),
    "Gestion des utilisateurs": (CAP3 + r"\fig08_utilisateurs.png", 15.5),
    "Création d'une demande": (CAP3 + r"\fig09_creation_mission.png", 15.5),
    "Circuit de validation": (CAP3 + r"\fig10_validations.png", 15.5),
    "Prise en charge logistique": (CAP3 + r"\fig11_dml_logistique.png", 15.5),
    "conventions hôtelières": (CAP3 + r"\fig12_conventions.png", 15.5),
    "parc véhicules": (CAP3 + r"\fig13_vehicules.png", 15.5),
    "Page Statistiques": (CAP3 + r"\fig14_statistiques.png", 15.5),
    "Messagerie interne": (CAP3 + r"\fig15_messagerie.png", 15.5),
    "Organigramme interactif": (CAP3 + r"\fig16_organigramme.png", 15.5),
    "diagramme de classes": (DIAG + r"\diagramme_classes.png", 15.0),
}

def renum(text):
    text = re.sub(r"Figure N°(\d+)", lambda m: "Figure N°%d" % (int(m.group(1)) + 3), text)
    text = re.sub(r"Tableau N°(\d+)", lambda m: "Tableau N°%d" % (int(m.group(1)) + 4), text)
    return text

inserted = 0
for el in list(src.element.body):
    tag = el.tag.split("}")[1]
    if tag == "sectPr":
        continue
    if tag == "p":
        text = "".join(n.text or "" for n in el.iter(qn("w:t"))).strip()
        if text.startswith("[Insérer ici"):
            hit = next((v for k, v in img_map.items() if k in text), None)
            if hit:
                image(hit[0], hit[1])
                inserted += 1
            else:
                print("PLACEHOLDER SANS IMAGE:", text[:70])
            continue
    new_el = copy.deepcopy(el)
    # renumerotation des legendes dans l'element copie
    for tnode in new_el.iter(qn("w:t")):
        if tnode.text and ("Figure N°" in tnode.text or "Tableau N°" in tnode.text):
            tnode.text = renum(tnode.text)
    BODY_SECT.addprevious(new_el)

section_break('rId101')   # fin du chapitre III -> en-tete Chapitre III
print("chapitre III OK, images:", inserted)

# ══════════════ 5. CHAPITRE IV ══════════════
para("CHAPITRE IV", style="Heading 1")
para("Réalisation", style="Heading 2")
para("Introduction", style="Heading 3")
para("Après la phase de conception présentée au chapitre précédent, ce chapitre décrit la phase de réalisation de l'application « AT Réservations » : les langages, bibliothèques et frameworks retenus, les logiciels utilisés durant le développement, la structure physique de la base de données, ainsi que les tests effectués pour valider le bon fonctionnement de la solution.")

para("Langages, bibliothèques et frameworks utilisés", style="Heading 3")
para("JavaScript (ES6+)", style="Heading 4")
para("JavaScript est le langage de programmation du web côté client. La syntaxe moderne ES6+ (modules, fonctions fléchées, déstructuration, async/await) a été utilisée pour écrire un code lisible et maintenable. Il a été choisi car il est le socle de l'écosystème React et bénéficie d'une communauté très active.")
para("PHP 8.2", style="Heading 4")
para("PHP est un langage de script côté serveur largement répandu. La version 8.2 apporte un typage renforcé et de meilleures performances. Il a été retenu car c'est le langage du framework Laravel et qu'il est parfaitement supporté par l'infrastructure d'hébergement interne (XAMPP/Apache).")
para("Dart", style="Heading 4")
para("Dart est le langage de programmation développé par Google pour le framework Flutter. Typé et orienté objet, il permet de compiler du code natif performant pour Android et iOS. Il a été choisi comme langage de l'application mobile AT Réservations.")
para("React 18", style="Heading 4")
para("React est une bibliothèque JavaScript de construction d'interfaces utilisateur basée sur des composants réutilisables. Les hooks (useState, useEffect, useMemo…) et le Context API assurent la gestion de l'état de l'application. React a été choisi pour sa performance (DOM virtuel), sa modularité et sa popularité en entreprise.")
para("Vite", style="Heading 4")
para("Vite est un outil de build et un serveur de développement nouvelle génération. Il offre un démarrage quasi instantané et un rechargement à chaud (HMR) très rapide, ce qui accélère considérablement le cycle de développement par rapport aux bundlers classiques.")
para("TailwindCSS", style="Heading 4")
para("TailwindCSS est un framework CSS « utility-first » : le style est composé directement dans le balisage à l'aide de classes utilitaires. Il a permis d'appliquer la charte graphique d'Algérie Télécom (vert #00A650, bleu #003DA5) de manière cohérente et de produire une interface entièrement responsive.")
para("Framer Motion", style="Heading 4")
para("Framer Motion est une bibliothèque d'animations pour React. Elle a été utilisée pour animer les transitions de pages, les cartes statistiques et les listes, afin d'offrir une expérience utilisateur fluide et moderne.")
para("Laravel 12", style="Heading 4")
para("Laravel est un framework PHP suivant le patron MVC (Modèle-Vue-Contrôleur). Il fournit l'ORM Eloquent, le système de migrations, la validation des requêtes, les politiques d'autorisation (Policies) et l'écosystème Artisan. Il constitue le cœur de l'API REST du backend : routage, contrôleurs, services métier, notifications email et journal d'audit.")
para("Laravel Sanctum", style="Heading 4")
para("Sanctum est le module d'authentification API de Laravel. Il délivre des jetons (tokens) révocables aux clients web et mobile, transportés dans des cookies httpOnly pour l'application web afin de renforcer la sécurité contre les attaques XSS.")
para("MySQL", style="Heading 4")
para("MySQL est le système de gestion de bases de données relationnelles retenu. Hébergé localement via XAMPP, il stocke l'ensemble des données de l'application (utilisateurs, missions, hôtels, conventions, véhicules, notifications, messages, journal d'audit) et garantit l'intégrité référentielle par des clés étrangères.")
para("Flutter", style="Heading 4")
para("Flutter est le framework mobile multiplateforme de Google. L'application mobile AT Réservations, développée en Flutter/Dart, permet aux utilisateurs nomades de consulter leurs missions, suivre les validations, recevoir les notifications et échanger des messages depuis un terminal Android ou iOS, en communiquant avec la même API REST que l'application web.")

para("Logiciels utilisés", style="Heading 3")
bullets([
    "Visual Studio Code : éditeur de code principal, avec extensions PHP, React et Flutter ;",
    "XAMPP : environnement serveur local regroupant Apache, MySQL et phpMyAdmin ;",
    "Postman : outil de test manuel des points d'accès (endpoints) de l'API REST ;",
    "GitHub : hébergement du dépôt Git, gestion de versions et intégration continue (CI) exécutant automatiquement les tests ;",
    "Claude Code CLI : assistant de développement par intelligence artificielle en ligne de commande, utilisé pour accélérer l'écriture du code, des tests et le diagnostic des anomalies ;",
    "Android Studio (émulateur) : exécution et test de l'application mobile Flutter.",
])

para("Les tables de la base de données MySQL", style="Heading 3")
para("Le Modèle Logique de Données Relationnel présenté au chapitre III a été implémenté dans MySQL au moyen des migrations Laravel. Le tableau suivant récapitule les principales tables de la base de données :")
kv_table([
    ("Table", "Rôle", "Colonnes principales"),
    ("users", "Comptes utilisateurs et rôles", "id, nom, prenom, email, password (bcrypt), role, direction, statut"),
    ("missions", "Demandes et ordres de mission", "id, titre, destination, date_depart, date_retour, type_mission, statut, motif_rejet, user_id, validateur_id"),
    ("hotels", "Référentiel des hôtels", "id, nom, ville, wilaya, etoiles, statut_convention"),
    ("conventions", "Classe d'association Mission–Hôtel", "id, mission_id, hotel_id, date_debut, date_fin, tarif_journalier, nombre_nuits"),
    ("vehicules", "Parc de véhicules de service", "id, marque, modele, immatriculation, statut, mission_id"),
    ("bons_transport", "Bons de transport émis", "id, numero, mission_id, type_transport, date_emission"),
    ("pieces_justificatives", "Pièces jointes des missions", "id, mission_id, nom_fichier, chemin_fichier, type"),
    ("notifications", "Notifications internes", "id, user_id, titre, contenu, lu, type"),
    ("messages", "Messagerie interne", "id, expediteur_id, destinataire_id, mission_id, contenu, lu"),
    ("audit_logs", "Journal d'audit des actions", "id, user_id, action, description, adresse_ip, created_at"),
])
caption("Tableau N°7 : Les tables de la base de données MySQL")

para("Les tests", style="Heading 3")
para("Test du design", style="Heading 4")
para("La conformité visuelle de l'application a été vérifiée page par page : respect des couleurs de la charte Algérie Télécom (vert #00A650 et bleu #003DA5), cohérence de la typographie, affichage correct en mode sombre, et adaptation responsive de l'interface aux différentes tailles d'écran (poste de travail, tablette, mobile).")
caption("Figure N°21 : [À compléter — capture du test responsive]")
para("Test sur les navigateurs", style="Heading 4")
para("L'application web a été testée sur les trois navigateurs exigés par le cahier des charges : Google Chrome, Mozilla Firefox et Microsoft Edge. Le rendu des pages, les formulaires, les exports PDF/Excel et les notifications fonctionnent de manière identique sur les trois navigateurs.")
caption("Figure N°22 : [À compléter — capture du test multi-navigateurs]")
para("Test des fonctionnalités", style="Heading 4")
para("Chaque fonctionnalité majeure a fait l'objet de tests fonctionnels de bout en bout : authentification (connexion, déconnexion, restrictions par rôle), création, modification et suppression de missions, circuit complet de validation (approbation et rejet avec motif), traitement logistique DML (affectation hôtel et véhicule), messagerie interne, notifications et exports PDF/Excel. En complément, une suite de 74 tests automatisés PHPUnit (182 assertions) couvre l'API backend : authentification, workflow de mission, politiques d'autorisation, en-têtes de sécurité et points d'accès de messagerie et de notifications. Ces tests sont exécutés automatiquement à chaque mise à jour du dépôt via GitHub Actions.")
caption("Figure N°23 : [À compléter — capture de l'exécution des tests PHPUnit]")

para("Conclusion", style="Heading 3")
para("Ce chapitre a présenté l'environnement technique de la réalisation : les langages et frameworks (React, Laravel, Flutter), les logiciels de développement, l'implémentation de la base de données MySQL ainsi que les différents tests menés pour garantir la qualité de l'application. L'ensemble de ces choix a permis d'aboutir à une solution complète, fonctionnelle et conforme aux besoins exprimés.")

section_break('rId102')   # fin du chapitre IV -> en-tete Chapitre IV
print("chapitre IV OK")

# ══════════════ 6. CONCLUSION GENERALE + BIBLIOGRAPHIE ══════════════
para("Conclusion Générale", style="Heading 3")
para("Au terme de ce projet de fin de formation, nous avons conçu et réalisé « AT Réservations », une application web intranet et mobile de gestion automatisée des missions et déplacements professionnels au profit d'Algérie Télécom. La solution couvre l'intégralité du processus métier : création des demandes par les agents, validation hiérarchique par les directeurs, prise en charge logistique par la Direction des Moyens Logistiques (hôtels en convention, parc véhicules, bons de transport), supervision et statistiques pour l'administration, le tout complété par une messagerie interne, des notifications automatiques par email et un journal d'audit garantissant la traçabilité.")
para("La réalisation de ce projet nous a confrontés à plusieurs difficultés : la modélisation d'un circuit de validation hiérarchique fidèle à l'organisation réelle de l'entreprise, la gestion fine des rôles et permissions, la sécurisation de l'authentification (jetons révocables, cookies httpOnly, hachage bcrypt), ainsi que la coordination entre trois couches techniques distinctes (API Laravel, client web React, application mobile Flutter). Ces difficultés ont été surmontées grâce à une démarche itérative faite de tests systématiques et de corrections successives.")
para("Ce travail ouvre plusieurs perspectives d'amélioration : l'intégration effective de l'annuaire Active Directory/LDAP de l'entreprise en environnement de production, l'enrichissement de l'application mobile Flutter (mode hors ligne, notifications push), la mise en place de notifications en temps réel par WebSockets, un tableau de bord décisionnel avancé avec analyse prédictive des budgets, ainsi qu'une passerelle vers le système comptable de l'entreprise pour le rapprochement automatique des dépenses de mission.")

para("Bibliographie / Webographie", style="Heading 3")
bullets([
    "[1] Site officiel d'Algérie Télécom — https://www.algerietelecom.dz",
    "[2] Loi n° 2000-03 du 5 août 2000 fixant les règles générales relatives à la poste et aux télécommunications — Journal Officiel de la République Algérienne",
    "[3] Documentation officielle React — https://react.dev",
    "[4] Documentation officielle Laravel — https://laravel.com/docs",
    "[5] Documentation officielle TailwindCSS — https://tailwindcss.com",
    "[6] Documentation officielle Flutter — https://flutter.dev",
    "[7] Documentation officielle MySQL — https://dev.mysql.com/doc",
    "[8] P.-A. Muller, N. Gaertner, « Modélisation objet avec UML », Éditions Eyrolles",
])
print("conclusion + biblio OK")

# ══════════════ 7. docPr uniques (corps) ══════════════
nid = 100
for docPr in d.element.body.iter(qn("wp:docPr")):
    docPr.set("id", str(nid))
    docPr.set("name", "Image %d" % nid)
    nid += 1

d.save(BASE)
print("MASTER assemble, images corps:", nid - 100)
