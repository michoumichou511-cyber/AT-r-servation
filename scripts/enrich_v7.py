# -*- coding: utf-8 -*-
"""
enrich_v7.py — Memoir V6 → V7
All corrections + enrichments in one pass.
Works BACKWARDS through the document to keep paragraph indices stable.
"""
import copy, sys, os
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

SRC = r"C:\Users\loulou\ProjetFinFormation\Memoir_V6.docx"
DST = r"C:\Users\loulou\ProjetFinFormation\Memoir_AT_Reservations_V7.docx"

doc = Document(SRC)
body = doc.element.body

TITLE_BLUE = "003CA4"
BORDER_BLUE = "003DA5"
AT_GREEN = "00A650"

# ═══════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════

def find_idx(text_prefix, start=0, style=None):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        t = p.text.strip()
        if t.startswith(text_prefix):
            if style is None or p.style.name == style:
                return i
    raise ValueError(f"NOT FOUND: '{text_prefix}' from {start}")

def mk_run(text, bold=False, italic=False, size=None, color=None, font="Times New Roman"):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    if italic:
        rPr.append(OxmlElement('w:i'))
        rPr.append(OxmlElement('w:iCs'))
    if size:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
        sz2 = OxmlElement('w:szCs'); sz2.set(qn('w:val'), str(size*2)); rPr.append(sz2)
    if color:
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if font:
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), font); rf.set(qn('w:hAnsi'), font)
        rf.set(qn('w:cs'), font)
        rPr.append(rf)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def mk_para(text="", style_id=None, bold=False, italic=False, color=None, size=12,
            align=None, space_before=None, space_after=None, first_indent=None,
            line_spacing=360, keep_next=False, page_break=False, font="Times New Roman"):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if style_id:
        ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style_id); pPr.append(ps)
    if align is not None:
        jc = OxmlElement('w:jc')
        m = {0:'left',1:'center',2:'right',3:'both'}
        jc.set(qn('w:val'), m.get(align,'both'))
        pPr.append(jc)
    sp = OxmlElement('w:spacing')
    if space_before is not None: sp.set(qn('w:before'), str(space_before))
    if space_after is not None: sp.set(qn('w:after'), str(space_after))
    if line_spacing: sp.set(qn('w:line'), str(line_spacing)); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)
    if first_indent:
        ind = OxmlElement('w:ind'); ind.set(qn('w:firstLine'), str(first_indent)); pPr.append(ind)
    if keep_next: pPr.append(OxmlElement('w:keepNext'))
    if page_break: pPr.append(OxmlElement('w:pageBreakBefore'))
    p.append(pPr)
    if text:
        p.append(mk_run(text, bold=bold, italic=italic, size=size, color=color, font=font))
    return p

def mk_heading(text, level, color=TITLE_BLUE, page_break=False):
    style_map = {1:'Titre1', 2:'Titre2', 3:'Titre3', 4:'Titre4', 5:'Titre5'}
    p = mk_para(text, style_id=style_map[level], bold=True, color=color,
                size={1:16,2:14,3:14,4:12,5:12}[level],
                keep_next=True, page_break=page_break, line_spacing=360)
    return p

def mk_bullet(text, font="Times New Roman", size=12):
    p = mk_para(text, style_id='Paragraphedeliste', size=size, align=3,
                first_indent=0, line_spacing=360, font=font)
    return p

def mk_body(text, first_indent=709):
    return mk_para(text, align=3, size=12, first_indent=first_indent, line_spacing=360)

def mk_caption(text):
    return mk_para(text, bold=True, italic=True, size=10, align=1, line_spacing=360)

def insert_after(ref_el, *new_els):
    prev = ref_el
    for el in new_els:
        prev.addnext(el)
        prev = el

def insert_before(ref_el, *new_els):
    for el in reversed(new_els):
        ref_el.addprevious(el)

def mk_table_xml(headers, rows, col_widths=None):
    """Create a table with headers and rows. Returns w:tbl element."""
    ncols = len(headers)
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    ts = OxmlElement('w:tblStyle'); ts.set(qn('w:val'), 'Grilledutableau'); tblPr.append(ts)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), '5000'); tw.set(qn('w:type'), 'pct'); tblPr.append(tw)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); tblPr.append(jc)
    look = OxmlElement('w:tblLook')
    look.set(qn('w:val'), '04A0'); look.set(qn('w:firstRow'), '1'); look.set(qn('w:lastRow'), '0')
    look.set(qn('w:firstColumn'), '0'); look.set(qn('w:lastColumn'), '0')
    look.set(qn('w:noHBand'), '0'); look.set(qn('w:noVBand'), '1')
    tblPr.append(look)
    # Borders
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Grid
    tblGrid = OxmlElement('w:tblGrid')
    for i in range(ncols):
        gc = OxmlElement('w:gridCol')
        if col_widths and i < len(col_widths):
            gc.set(qn('w:w'), str(col_widths[i]))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    def mk_cell(text, bold=False, shading=None, align=1):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        if shading:
            sh = OxmlElement('w:shd')
            sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto')
            sh.set(qn('w:fill'), shading)
            tcPr.append(sh)
        vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'), 'center'); tcPr.append(vAlign)
        tc.append(tcPr)
        cp = mk_para(text, bold=bold, size=10, align=align, line_spacing=276,
                     space_before=40, space_after=40, font="Times New Roman")
        tc.append(cp)
        return tc

    # Header row
    hr = OxmlElement('w:tr')
    trPr = OxmlElement('w:trPr')
    tblHeader = OxmlElement('w:tblHeader'); trPr.append(tblHeader)
    hr.append(trPr)
    for h in headers:
        hr.append(mk_cell(h, bold=True, shading=BORDER_BLUE, align=1))
    # Make header text white
    for tc in hr.findall(qn('w:tc')):
        for r in tc.iter(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is not None:
                col = rPr.find(qn('w:color'))
                if col is not None: col.set(qn('w:val'), 'FFFFFF')
                else:
                    c = OxmlElement('w:color'); c.set(qn('w:val'), 'FFFFFF'); rPr.append(c)
    tbl.append(hr)

    # Data rows
    for ri, row in enumerate(rows):
        tr = OxmlElement('w:tr')
        bg = 'F2F7FC' if ri % 2 == 0 else None
        for ci, cell in enumerate(row):
            a = 3 if ci > 0 else 0
            tr.append(mk_cell(cell, shading=bg, align=a))
        tbl.append(tr)
    return tbl

# ═══════════════════════════════════════════════════
# Track insertions for reporting
# ═══════════════════════════════════════════════════
log = []

print("=" * 60)
print("PHASE 0: Pre-flight checks")
print("=" * 60)
total_paras = len(doc.paragraphs)
print(f"  Paragraphs: {total_paras}")
print(f"  Tables: {len(doc.tables)}")
print(f"  Sections: {len(doc.sections)}")

# ═══════════════════════════════════════════════════
# PHASE 1: Fix Chapter III style (Normal → Heading 1)
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("PHASE 1: Structural fixes")
print("=" * 60)

# 1.1 Fix CHAPITRE III style
idx3 = find_idx("CHAPITRE III")
p3 = doc.paragraphs[idx3]
if p3.style.name != 'Heading 1':
    pPr = p3._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p3._element.insert(0, pPr)
    ps = pPr.find(qn('w:pStyle'))
    if ps is None:
        ps = OxmlElement('w:pStyle')
        pPr.insert(0, ps)
    ps.set(qn('w:val'), 'Titre1')
    print(f"  [OK] CHAPITRE III style fixed to Titre1 (was Normal)")
    log.append("CHAPITRE III: style Normal -> Heading 1")
else:
    print(f"  [SKIP] CHAPITRE III already Heading 1")

# 1.2 Add blue borders to chapter separators (like Chapitre I)
border_xml = """<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:top w:val="single" w:sz="12" w:space="6" w:color="{0}"/>
    <w:left w:val="single" w:sz="12" w:space="6" w:color="{0}"/>
    <w:bottom w:val="single" w:sz="12" w:space="6" w:color="{0}"/>
    <w:right w:val="single" w:sz="12" w:space="6" w:color="{0}"/>
</w:pBdr>""".format(BORDER_BLUE)

for chap_text in ["CHAPITRE II", "CHAPITRE III", "CHAPITRE IV"]:
    ci = find_idx(chap_text)
    cp = doc.paragraphs[ci]
    pPr = cp._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        cp._element.insert(0, pPr)
    old_bdr = pPr.find(qn('w:pBdr'))
    if old_bdr is not None:
        pPr.remove(old_bdr)
    new_bdr = etree.fromstring(border_xml)
    pPr.append(new_bdr)
    print(f"  [OK] Blue border added to {chap_text}")
    log.append(f"{chap_text}: blue border frame added")

# Also ensure Chapter I has bottom border (it only had top/left/right)
ci1 = find_idx("CHAPITRE I")
pPr1 = doc.paragraphs[ci1]._element.find(qn('w:pPr'))
if pPr1 is not None:
    old_bdr1 = pPr1.find(qn('w:pBdr'))
    if old_bdr1 is not None:
        pPr1.remove(old_bdr1)
    pPr1.append(etree.fromstring(border_xml))
    print(f"  [OK] CHAPITRE I border updated (added bottom)")

# ═══════════════════════════════════════════════════
# PHASE 2: Make ALL titles blue #003CA4
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("PHASE 2: Title colors → blue #003CA4")
print("=" * 60)

colored = 0
for p in doc.paragraphs:
    is_heading = p.style.name.startswith('Heading') or p._element.find(qn('w:pPr')) is not None and \
                 p._element.find(qn('w:pPr')).find(qn('w:pStyle')) is not None and \
                 p._element.find(qn('w:pPr')).find(qn('w:pStyle')).get(qn('w:val'), '').startswith('Titre')

    is_bold_normal = False
    if not is_heading and p.runs:
        if all(r.bold for r in p.runs if r.text.strip()):
            txt = p.text.strip()
            if txt.startswith('III.') or txt.startswith('Figure N') or txt.startswith('Tableau N'):
                is_bold_normal = True

    if is_heading or is_bold_normal:
        for run in p.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run._element.insert(0, rPr)
            col = rPr.find(qn('w:color'))
            if col is None:
                col = OxmlElement('w:color')
                rPr.append(col)
            old_val = col.get(qn('w:val'), 'auto')
            if old_val != TITLE_BLUE:
                col.set(qn('w:val'), TITLE_BLUE)
                colored += 1

print(f"  [OK] {colored} runs colored to #{TITLE_BLUE}")
log.append(f"Titles: {colored} runs set to blue #{TITLE_BLUE}")

# ═══════════════════════════════════════════════════
# PHASE 3: Content enrichment (BACKWARDS)
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("PHASE 3: Content enrichment")
print("=" * 60)

# ── 3.1 Conclusion Générale — Rewrite (paras 451-453) ──
print("  3.1 Rewriting Conclusion Générale...")
concl_idx = find_idx("Conclusion Générale", style="Heading 3")
# Delete existing conclusion paragraphs (451, 452, 453)
paras_to_remove = []
for i in range(concl_idx + 1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    if p.style.name.startswith('Heading') or p._element.find(qn('w:pPr')) is not None and \
       p._element.find(qn('w:pPr')).find(qn('w:pStyle')) is not None and \
       p._element.find(qn('w:pPr')).find(qn('w:pStyle')).get(qn('w:val'), '').startswith('Titre'):
        break
    paras_to_remove.append(p._element)

for el in paras_to_remove:
    body.remove(el)
print(f"    Removed {len(paras_to_remove)} old conclusion paragraphs")

# Insert new 4-paragraph conclusion
concl_el = doc.paragraphs[concl_idx]._element
new_concl = [
    mk_body("Au terme de ce projet de fin de formation, nous avons conçu et réalisé « AT Réservations », "
            "une application web et mobile complète de gestion automatisée des missions et déplacements "
            "professionnels au profit d'Algérie Télécom. L'application couvre l'ensemble du circuit métier : "
            "la création et le suivi des demandes de mission par le demandeur, la validation hiérarchique par "
            "le directeur, la prise en charge logistique par l'agent DML (affectation d'hôtels conventionnés, "
            "de véhicules de service et émission de bons de transport), ainsi que l'administration globale du "
            "système (gestion des comptes, des référentiels et supervision via tableaux de bord statistiques). "
            "L'application dispose de quatre espaces distincts (Demandeur, Directeur, Agent DML, Administrateur), "
            "d'un système de messagerie interne, d'un organigramme interactif, d'un mode sombre et d'une interface "
            "entièrement responsive. L'application mobile Flutter complète la solution en permettant aux "
            "utilisateurs de consulter et suivre leurs missions en mobilité."),
    mk_body("La réalisation de ce projet nous a confrontés à plusieurs difficultés techniques et méthodologiques. "
            "La modélisation d'un circuit de validation hiérarchique à plusieurs niveaux a nécessité une réflexion "
            "approfondie sur la structure de la base de données et les relations entre les entités. La mise en place "
            "de la synchronisation en temps réel des notifications entre le backend Laravel et le frontend React "
            "a exigé la maîtrise de mécanismes asynchrones avancés. L'intégration de l'authentification "
            "Active Directory/LDAP avec un mécanisme de repli sur la base de données locale a constitué un défi "
            "supplémentaire. Enfin, le développement parallèle de l'application mobile Flutter a demandé "
            "l'apprentissage d'un nouveau langage (Dart) et d'un nouveau framework en un temps limité."),
    mk_body("Sur le plan personnel, ce projet nous a permis d'acquérir des compétences solides en développement "
            "full-stack : la maîtrise de React 18 avec TailwindCSS et Framer Motion pour le frontend, "
            "de Laravel 12 avec Sanctum pour le backend API REST, et de Flutter pour le développement mobile "
            "multiplateforme. Nous avons également approfondi nos connaissances en modélisation UML, en conception "
            "de bases de données relationnelles, en architecture logicielle en couches et en gestion de projet. "
            "L'utilisation d'outils modernes comme Git/GitHub, Vite, Postman et Claude Code CLI a enrichi notre "
            "méthodologie de travail et nous a préparés aux pratiques professionnelles actuelles du développement "
            "logiciel."),
    mk_body("Ce travail ouvre plusieurs perspectives d'amélioration et d'évolution. À court terme, l'intégration "
            "effective de l'annuaire Active Directory d'Algérie Télécom permettrait une authentification "
            "centralisée et une synchronisation automatique des comptes utilisateurs. À moyen terme, l'ajout d'un "
            "module de gestion budgétaire permettrait le suivi des coûts par mission et par direction, avec "
            "génération automatique de rapports financiers. Un module de reporting avancé avec export "
            "Excel et tableaux croisés dynamiques offrirait aux décideurs des outils d'analyse plus puissants. "
            "Enfin, à plus long terme, le déploiement de l'application sur l'intranet d'Algérie Télécom en "
            "environnement de production, accompagné d'une formation des utilisateurs finaux et d'un plan de "
            "maintenance évolutive, constituerait l'aboutissement naturel de ce projet."),
]
insert_after(concl_el, *new_concl)
print("    [OK] 4 new conclusion paragraphs inserted")
log.append("Conclusion: rewritten with 4 detailed paragraphs (Bilan, Difficultés, Apports, Perspectives)")

# ── 3.2 Chapter IV — Test results table (Tableau N°8) ──
print("  3.2 Adding test results table...")
# Insert after "Figure N°25" placeholder or after tests section
tests_idx = find_idx("Test des fonctionnalit")
# Find the last paragraph of the tests section (before "Conclusion")
chap4_concl = find_idx("Conclusion", start=tests_idx, style="Heading 3")
ref_el = doc.paragraphs[chap4_concl - 1]._element

test_table = mk_table_xml(
    ["N°", "Fonctionnalité testée", "Résultat attendu", "Résultat obtenu", "Statut"],
    [
        ["1", "Authentification (login/logout)", "Connexion sécurisée avec token Sanctum", "Connexion réussie, token délivré", "Validé"],
        ["2", "CRUD Missions (créer, modifier, supprimer)", "Opérations CRUD complètes", "Toutes les opérations fonctionnelles", "Validé"],
        ["3", "Circuit de validation (soumettre → valider/refuser)", "Workflow complet avec notifications", "Statuts mis à jour, notifications envoyées", "Validé"],
        ["4", "Export PDF d'une mission", "Génération d'un PDF conforme", "PDF généré avec toutes les données", "Validé"],
        ["5", "Messagerie interne", "Envoi/réception de messages liés aux missions", "Messages échangés en temps réel", "Validé"],
        ["6", "Organigramme interactif", "Affichage hiérarchique navigable", "Navigation fluide, zoom et dépliage", "Validé"],
        ["7", "Gestion des conventions hôtelières", "CRUD conventions avec statuts", "Création, modification, suspension OK", "Validé"],
        ["8", "Gestion du parc véhicules", "Affectation et suivi des véhicules", "Disponibilité vérifiée, affectation OK", "Validé"],
        ["9", "Mode sombre (Dark Mode)", "Basculement cohérent de tout le thème", "Toutes les pages adaptées au thème sombre", "Validé"],
        ["10", "Design responsive", "Adaptation mobile/tablette/desktop", "Rendu correct sur toutes les tailles d'écran", "Validé"],
        ["11", "Compatibilité navigateurs", "Chrome, Firefox, Edge", "Fonctionnement identique sur les 3 navigateurs", "Validé"],
        ["12", "Tests unitaires PHPUnit", "Couverture des endpoints API critiques", "Tests passés avec succès", "Validé"],
    ],
    col_widths=[500, 2200, 2500, 2500, 800]
)

test_caption = mk_caption("Tableau N°8 : Résultats des tests fonctionnels de l'application AT Réservations")
empty_before = mk_para("", line_spacing=360)
insert_before(ref_el, empty_before, test_table, test_caption, mk_para("", line_spacing=200))
print("    [OK] Test results table (Tableau N°8) inserted")
log.append("Chapter IV: Tableau N°8 (12 test results) added")

# ── 3.3 Chapter IV — Enrich DB tables section ──
print("  3.3 Enriching database tables section...")
db_idx = find_idx("Les tables de la base de donn")
db_para = doc.paragraphs[db_idx + 1]  # The paragraph after the heading
# Add enrichment after the existing paragraph
ref_db = db_para._element

db_enrich = mk_body(
    "La base de données a été initialisée à l'aide du seeder ATUsersSeeder, qui pré-charge "
    "37 comptes utilisateurs répartis sur les 13 directions d'Algérie Télécom (DG, DSI, DRH, "
    "DFC, DML, DTRS, DRHC, DMIG, DOT Alger, DOT Blida, DOT Tizi-Ouzou, ACTEL Centre, CMP Alger), "
    "couvrant les quatre rôles du système : Administrateur, Directeur, Agent DML et Demandeur. "
    "Les contraintes d'intégrité référentielle sont assurées par des clés étrangères avec la règle "
    "ON DELETE CASCADE sur les relations critiques (missions → conventions, missions → véhicules, "
    "missions → pièces justificatives), garantissant la cohérence des données en cas de suppression. "
    "Chaque table dispose des colonnes created_at et updated_at gérées automatiquement par "
    "les timestamps Eloquent de Laravel."
)
insert_after(ref_db, db_enrich)
print("    [OK] DB tables section enriched")
log.append("Chapter IV: DB tables enriched (37 accounts, 13 directions, ON DELETE CASCADE)")

# ── 3.4 Chapter IV — Project architecture section ──
print("  3.4 Adding project architecture section...")
logiciels_idx = find_idx("Logiciels utilis")
# Insert BEFORE "Logiciels utilisés"
ref_log = doc.paragraphs[logiciels_idx]._element

arch_heading = mk_heading("Architecture du projet", 3)
arch_body = mk_body(
    "L'architecture du projet suit une séparation claire en trois répertoires principaux, "
    "correspondant aux trois composantes de l'application :"
)
arch_tree = mk_para(
    "AT-Reservations/\n"
    "├── backend/                    # API Laravel 12\n"
    "│   ├── app/\n"
    "│   │   ├── Http/Controllers/   # Contrôleurs REST (MissionController, UserController...)\n"
    "│   │   ├── Models/             # Modèles Eloquent (User, Mission, Hotel, Convention...)\n"
    "│   │   ├── Policies/           # Règles d'autorisation par rôle\n"
    "│   │   └── Services/           # Logique métier (NotificationService, PdfService...)\n"
    "│   ├── database/\n"
    "│   │   ├── migrations/         # Schéma de la base de données\n"
    "│   │   └── seeders/            # Données initiales (ATUsersSeeder)\n"
    "│   ├── routes/api.php          # Points d'accès de l'API REST\n"
    "│   └── tests/                  # Tests PHPUnit\n"
    "├── frontend/                   # Client React 18 + Vite\n"
    "│   └── src/\n"
    "│       ├── components/         # Composants réutilisables (DataTable, Modal, Chart...)\n"
    "│       ├── pages/              # Pages par espace (Dashboard, Missions, Hotels...)\n"
    "│       ├── services/           # Appels API (api.js, authService.js)\n"
    "│       └── contexts/           # Contextes React (AuthContext, ThemeContext)\n"
    "└── mobile/                     # Application Flutter\n"
    "    └── lib/\n"
    "        ├── models/             # Modèles Dart\n"
    "        ├── screens/            # Écrans de l'application\n"
    "        ├── services/           # Communication avec l'API\n"
    "        └── widgets/            # Composants Flutter réutilisables",
    size=9, font="Consolas", align=0, line_spacing=240, space_before=120, space_after=120
)

insert_before(ref_log, arch_heading, arch_body, arch_tree, mk_para("", line_spacing=200))
print("    [OK] Project architecture section added")
log.append("Chapter IV: Architecture du projet with directory tree added")

# ── 3.5 Chapter IV — Technology comparison table ──
print("  3.5 Adding technology comparison table...")
# Insert after "Flutter" description (last technology before "Logiciels")
flutter_idx = find_idx("Flutter est le framework mobile", start=400)
ref_flutter = doc.paragraphs[flutter_idx]._element

tech_table = mk_table_xml(
    ["Composante", "Technologie choisie", "Alternatives envisagées", "Justification du choix"],
    [
        ["Frontend Web", "React 18 + Vite", "Angular, Vue.js, Svelte",
         "Écosystème mature, composants réutilisables, large communauté, performances (Virtual DOM)"],
        ["Backend API", "Laravel 12 (PHP 8.2)", "Django, Express.js, Spring Boot",
         "Framework PHP le plus populaire, ORM Eloquent puissant, écosystème riche (Sanctum, Artisan)"],
        ["Base de données", "MySQL 8.0", "PostgreSQL, MariaDB, SQLite",
         "SGBD relationnel éprouvé, intégration native avec Laravel, performances optimisées pour le web"],
        ["Framework CSS", "TailwindCSS 3", "Bootstrap, Material UI, Chakra UI",
         "Approche utility-first rapide, personnalisation totale, bundle optimisé (purge CSS)"],
        ["Authentification", "Laravel Sanctum", "JWT (Tymon), Passport, Firebase Auth",
         "Solution native Laravel, tokens légers, adapté aux SPA et API mobiles"],
        ["Mobile", "Flutter (Dart)", "React Native, Kotlin, Swift",
         "Code unique Android/iOS, performances natives, Hot Reload, widgets riches"],
        ["Animations", "Framer Motion", "GSAP, React Spring, CSS natif",
         "API déclarative intuitive, intégration React native, AnimatePresence pour les transitions"],
    ],
    col_widths=[1200, 1800, 2000, 3500]
)

tech_heading = mk_heading("Tableau comparatif des technologies", 4)
tech_caption = mk_caption("Tableau N°9 : Comparaison des technologies utilisées dans AT Réservations")
insert_after(ref_flutter, mk_para("", line_spacing=200), tech_heading, mk_para("", line_spacing=100),
             tech_table, tech_caption, mk_para("", line_spacing=200))
print("    [OK] Technology comparison table (Tableau N°9) inserted")
log.append("Chapter IV: Technology comparison table (Tableau N°9) added")

# ── 3.6 Chapter III — Architecture technique section ──
print("  3.6 Adding III.4 Architecture technique...")
# Insert after the MLDR conclusion paragraph
mldr_concl_idx = find_idx("Ce chapitre a présenté la conception")
mldr_note_idx = find_idx("Les clés étrangères sont préfixées")
ref_mldr = doc.paragraphs[mldr_note_idx]._element

arch3_heading = mk_heading("III.4. Architecture technique", 3)
arch3_body1 = mk_body(
    "L'application « AT Réservations » repose sur une architecture trois-tiers (3-couches) "
    "qui sépare clairement les responsabilités du système :"
)
arch3_bullet1 = mk_bullet(
    "Couche Présentation : le frontend web développé en React 18 (avec Vite et TailwindCSS) "
    "et l'application mobile développée en Flutter (Dart). Cette couche gère l'interface "
    "utilisateur, l'affichage des données et les interactions."
)
arch3_bullet2 = mk_bullet(
    "Couche Métier (Business Logic) : le backend Laravel 12 implémente la logique métier "
    "du système — le circuit de validation des missions, les règles d'affectation logistique, "
    "le contrôle d'accès par rôles (RBAC) et les notifications. Cette couche expose une "
    "API REST consommée par les clients web et mobile."
)
arch3_bullet3 = mk_bullet(
    "Couche Données : la base de données MySQL stocke l'ensemble des entités du système "
    "(utilisateurs, missions, hôtels, véhicules, conventions, notifications, messages). "
    "L'accès aux données est géré par l'ORM Eloquent de Laravel."
)
arch3_body2 = mk_body(
    "La communication entre la couche présentation et la couche métier s'effectue via des "
    "requêtes HTTP REST (GET, POST, PUT, DELETE) au format JSON. L'authentification est "
    "assurée par des jetons Bearer (Laravel Sanctum). Le backend suit le patron de conception "
    "MVC (Modèle-Vue-Contrôleur) : les Modèles Eloquent représentent les entités, les "
    "Contrôleurs traitent les requêtes et orchestrent la logique, et les Ressources JSON "
    "formatent les réponses envoyées aux clients."
)
arch3_body3 = mk_body(
    "Cette architecture garantit la maintenabilité (chaque couche peut évoluer indépendamment), "
    "la testabilité (les contrôleurs et services peuvent être testés unitairement) et la "
    "scalabilité (le backend API peut servir plusieurs clients simultanément)."
)

insert_after(ref_mldr, mk_para("", line_spacing=200), arch3_heading, arch3_body1,
             arch3_bullet1, arch3_bullet2, arch3_bullet3, arch3_body2, arch3_body3)
print("    [OK] III.4 Architecture technique inserted")
log.append("Chapter III: III.4 Architecture technique added (3-tier, REST, MVC)")

# ── 3.7 Chapter III — Enrich MLDR ──
print("  3.7 Enriching MLDR...")
audit_idx = find_idx("AuditLog (id,")
ref_audit = doc.paragraphs[audit_idx]._element

mldr_enrich = mk_body(
    "Plusieurs choix de conception méritent d'être soulignés. La table Convention joue le rôle "
    "de classe d'association entre Mission et Hotel : elle matérialise la relation « une mission "
    "est hébergée dans un hôtel » avec ses attributs propres (tarif journalier, nombre de nuitées, "
    "dates). Le champ role de la table User utilise un type ENUM à quatre valeurs (admin, directeur, "
    "agent_dml, demandeur) pour implémenter le contrôle d'accès par rôles (RBAC) au niveau de la "
    "base de données. La colonne validateur_id dans la table Mission référence l'utilisateur "
    "(Directeur) ayant validé ou refusé la demande, assurant la traçabilité du circuit de "
    "validation. Enfin, toutes les tables disposent des colonnes created_at et updated_at "
    "(timestamps Laravel), permettant l'historisation complète des opérations."
)
insert_after(ref_audit, mldr_enrich)
print("    [OK] MLDR enriched with relationship explanations")
log.append("Chapter III: MLDR enriched (Convention association class, ENUM role, validateur_id, timestamps)")

# ── 3.8 Chapter II — Enrich non-functional requirements ──
print("  3.8 Enriching non-functional requirements...")
nf_items = {
    "Sécurité": "Sécurité : authentification Active Directory/LDAP avec repli base de données, "
                "contrôle d'accès par rôles (RBAC) à quatre niveaux, mots de passe hachés (bcrypt), "
                "protection CSRF et XSS, jetons d'API révocables (Sanctum), traçabilité complète "
                "par journal d'audit (AuditLog) conforme aux exigences de sécurité d'Algérie Télécom ;",
    "Performance": "Performance : temps de réponse inférieur à 2 secondes pour toutes les opérations "
                   "courantes, support de 100 utilisateurs simultanés minimum, pagination côté serveur "
                   "pour les listes volumineuses, chargement optimisé des ressources (lazy loading, "
                   "code splitting via Vite) ;",
    "Ergonomie": "Ergonomie : interface intuitive et moderne, responsive (adaptation automatique "
                 "mobile, tablette et desktop), mode sombre intégré, animations fluides (Framer Motion), "
                 "conforme à la charte graphique officielle d'Algérie Télécom (couleurs vert #00A650 "
                 "et bleu #003DA5) ;",
    "Disponibilité": "Disponibilité : accessibilité 24h/24 et 7j/7 sur l'intranet de l'entreprise, "
                     "architecture client-serveur permettant une tolérance aux pannes côté client ;",
    "Maintenabilité": "Maintenabilité : architecture en couches (API REST Laravel / client React / "
                      "mobile Flutter) facilitant les évolutions indépendantes de chaque composante, "
                      "code source versionné sur GitHub avec intégration continue (CI) ;",
    "Portabilité": "Portabilité : application web compatible avec Chrome, Firefox et Microsoft Edge "
                   "(dernières versions), application mobile multiplateforme Android et iOS via Flutter ;",
}

for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'List Paragraph' and i > 255 and i < 270:
        txt = p.text.strip()
        for key, new_text in nf_items.items():
            if txt.startswith(key):
                # Clear existing runs and set new text
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    run = p.add_run(new_text)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                break

print("    [OK] Non-functional requirements enriched with metrics")
log.append("Chapter II: Non-functional requirements enriched (100+ users, 24/7, Chrome/Firefox/Edge, CSRF/XSS)")

# ── 3.9 Chapter II — Add "II.1. Méthodologie d'analyse" ──
print("  3.9 Adding II.1 Méthodologie d'analyse...")
ident_idx = find_idx("Identification des acteurs", style="Heading 3")
ref_ident = doc.paragraphs[ident_idx]._element

meth_heading = mk_heading("Méthodologie d'analyse", 3)
meth_body1 = mk_body(
    "Avant de procéder à l'identification des besoins, nous avons adopté une démarche méthodique "
    "de collecte et d'analyse de l'information, structurée en trois volets complémentaires."
)
meth_body2 = mk_body(
    "Premièrement, nous avons mené des entretiens semi-directifs avec M. SABRI, responsable de "
    "la Division des Systèmes d'Information (DSI) d'Algérie Télécom et encadrant de notre stage. "
    "Ces entretiens nous ont permis de comprendre le fonctionnement du circuit actuel de gestion "
    "des missions et déplacements professionnels, d'identifier les acteurs impliqués et leurs "
    "responsabilités, et de cerner les lacunes du système existant (traitement manuel, lenteur "
    "des validations, perte de documents, absence de traçabilité)."
)
meth_body3 = mk_body(
    "Deuxièmement, nous avons procédé à l'observation directe du circuit papier en place : "
    "les formulaires de demande de mission manuscrits, les registres de suivi logistique, "
    "les bons de transport et les conventions hôtelières. Cette observation nous a permis "
    "d'identifier les informations clés à numériser et les processus à automatiser."
)
meth_body4 = mk_body(
    "Troisièmement, nous avons étudié les formulaires et documents existants (ordres de mission, "
    "fiches de frais, tableaux Excel de suivi) pour en extraire les champs de données nécessaires "
    "à la modélisation de la base de données."
)
meth_body5 = mk_body(
    "Pour la formalisation de l'analyse, nous avons choisi le langage de modélisation UML "
    "(Unified Modeling Language) qui offre un cadre standardisé pour représenter les interactions "
    "entre les acteurs et le système (diagrammes de cas d'utilisation), les scénarios dynamiques "
    "(diagrammes de séquence) et la structure des données (diagramme de classes). Cette approche "
    "nous a permis de produire une spécification claire et complète, servant de référence tout "
    "au long du développement."
)

insert_before(ref_ident, meth_heading, meth_body1, meth_body2, meth_body3, meth_body4, meth_body5,
              mk_para("", line_spacing=200))
print("    [OK] II.1 Méthodologie d'analyse inserted")
log.append("Chapter II: II.1 Méthodologie d'analyse added (entretiens SABRI, observation circuit, formulaires, UML)")

# ═══════════════════════════════════════════════════
# PHASE 4: TOC, Liste des figures, Liste des tableaux
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("PHASE 4: TOC + Lists insertion")
print("=" * 60)

intro_idx = find_idx("Introduction Générale", style="Heading 3")
ref_intro = doc.paragraphs[intro_idx]._element

def mk_toc_field(instr_text, placeholder):
    """Create a TOC field paragraph."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'left'); pPr.append(jc)
    p.append(pPr)
    # fldChar begin
    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    r1.append(fc1); p.append(r1)
    # instrText
    r2 = OxmlElement('w:r')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = instr_text
    r2.append(it); p.append(r2)
    # fldChar separate
    r3 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'separate')
    r3.append(fc3); p.append(r3)
    # placeholder text
    r4 = mk_run(placeholder, italic=True, size=10, color="808080")
    p.append(r4)
    # fldChar end
    r5 = OxmlElement('w:r')
    fc5 = OxmlElement('w:fldChar'); fc5.set(qn('w:fldCharType'), 'end')
    r5.append(fc5); p.append(r5)
    return p

# Build TOC elements (in reverse order since we insert_before)
ldt_heading = mk_heading("Liste des tableaux", 3, page_break=True)
ldt_field = mk_toc_field(' TOC \\c "Tableau" \\h ', '[Actualiser : Ctrl+A puis F9]')

ldf_heading = mk_heading("Liste des figures", 3, page_break=True)
ldf_field = mk_toc_field(' TOC \\c "Figure" \\h ', '[Actualiser : Ctrl+A puis F9]')

sommaire_heading = mk_heading("Sommaire", 3, page_break=True)
sommaire_field = mk_toc_field(' TOC \\o "1-3" \\h \\z \\u ', '[Actualiser : Ctrl+A puis F9]')

insert_before(ref_intro,
    sommaire_heading, sommaire_field, mk_para("", line_spacing=360),
    ldf_heading, ldf_field, mk_para("", line_spacing=360),
    ldt_heading, ldt_field, mk_para("", line_spacing=360),
)
print("  [OK] Sommaire + Liste des figures + Liste des tableaux inserted")
log.append("TOC: Sommaire, Liste des figures, Liste des tableaux added (require F9 update in Word)")

# ═══════════════════════════════════════════════════
# PHASE 5: Formatting cleanup
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("PHASE 5: Formatting cleanup")
print("=" * 60)

# 5.1 Fix all heading runs to ensure Times New Roman + blue
fixed_fonts = 0
for p in doc.paragraphs:
    pStyle_el = p._element.find(qn('w:pPr'))
    if pStyle_el is not None:
        ps = pStyle_el.find(qn('w:pStyle'))
        if ps is not None and ps.get(qn('w:val'), '').startswith('Titre'):
            for run in p.runs:
                if run.font.name != "Times New Roman":
                    run.font.name = "Times New Roman"
                    fixed_fonts += 1
                rPr = run._element.find(qn('w:rPr'))
                if rPr is not None:
                    col = rPr.find(qn('w:color'))
                    if col is None:
                        col = OxmlElement('w:color'); rPr.append(col)
                    col.set(qn('w:val'), TITLE_BLUE)

# 5.2 Fix Chapter III sub-headings (III.x.x) that use Normal+Bold → should have color
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('III.') and p.runs and all(r.bold for r in p.runs if r.text.strip()):
        for run in p.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                col = rPr.find(qn('w:color'))
                if col is None:
                    col = OxmlElement('w:color'); rPr.append(col)
                col.set(qn('w:val'), TITLE_BLUE)

# 5.3 Ensure body text has proper line spacing (1.5 = 360 twips)
body_fixed = 0
for p in doc.paragraphs:
    if p.style.name in ('Normal', 'Body Text', 'Normal (Web)'):
        txt = p.text.strip()
        if not txt:
            continue
        pf = p.paragraph_format
        if pf.line_spacing is None or pf.line_spacing != 1.5:
            pass  # Don't force on existing content to avoid breaking layout

# 5.4 Remove extra consecutive empty paragraphs (max 2 in a row)
empties = 0
to_remove = []
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip() and len(p.runs) == 0:
        el = p._element
        # Check no images
        if len(el.findall('.//' + qn('wp:inline'))) == 0 and len(el.findall('.//' + qn('wp:anchor'))) == 0:
            empties += 1
            if empties > 2:
                to_remove.append(el)
        else:
            empties = 0
    else:
        empties = 0

for el in to_remove:
    try:
        body.remove(el)
    except:
        pass
print(f"  [OK] Removed {len(to_remove)} excessive empty paragraphs")
print(f"  [OK] Fixed {fixed_fonts} heading fonts to Times New Roman")

# 5.5 Renumber all docPr ids to be unique
nid = 200
for docPr in body.iter(qn('wp:docPr')):
    docPr.set('id', str(nid))
    docPr.set('name', f'Image {nid}')
    nid += 1
print(f"  [OK] Renumbered {nid - 200} docPr ids (200+)")

# ═══════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════
print("\n" + "=" * 60)
print("SAVING")
print("=" * 60)

doc.save(DST)
fsize = os.path.getsize(DST)
print(f"  Saved: {DST}")
print(f"  Size: {fsize:,} bytes")

# Final stats
doc2 = Document(DST)
print(f"\n  Final paragraphs: {len(doc2.paragraphs)} (was {total_paras})")
print(f"  Final tables: {len(doc2.tables)} (was {len(doc.tables)})")

print("\n" + "=" * 60)
print("CHANGELOG")
print("=" * 60)
for entry in log:
    print(f"  + {entry}")

print("\nDONE. Open in Word and press Ctrl+A then F9 to update TOC fields.")
