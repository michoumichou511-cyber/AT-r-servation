from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "FONCTIONNALITES_PAGES_ADMIN_ET_APP.md"
OUT = ROOT / "docs" / "FONCTIONNALITES_PAGES_ADMIN_ET_APP.docx"


def add_markdown_like(document: Document, md_text: str) -> None:
    """
    Minimal MD -> DOCX conversion:
    - # / ## / ### headings
    - '- ' bullet points
    - paragraphs
    """
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        # quote-ish / note lines
        if line.startswith("> "):
            document.add_paragraph(line[2:].strip(), style="Intense Quote")
            continue

        # horizontal rules
        if set(line.strip()) == {"-"}:
            document.add_paragraph("")
            continue

        document.add_paragraph(line)


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Missing source file: {SRC}")

    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    add_markdown_like(doc, md)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK: wrote {OUT}")


if __name__ == "__main__":
    main()

