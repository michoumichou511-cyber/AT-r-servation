from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Couleurs AT ───────────────────────────────────────────────────────────────
AT_GREEN = RGBColor(0x00, 0xA6, 0x50)
AT_BLUE  = RGBColor(0x00, 0x3D, 0xA5)
BLACK    = RGBColor(0x00, 0x00, 0x00)
GREY     = RGBColor(0x60, 0x60, 0x60)

# ── Styles helpers ────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=BLACK, name='Times New Roman'):
    run.font.name        = name
    run.font.size        = Pt(size)
    run.font.bold        = bold
    run.font.italic      = italic
    run.font.color.rgb   = color

def add_paragraph(doc, text='', style='Normal', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=11, bold=False, italic=False, color=BLACK, space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size=16, bold=True, color=AT_BLUE)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=AT_BLUE)
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=AT_GREEN)
    return p

def add_heading4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, italic=True, color=BLACK)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p

def add_note(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(label)
    set_run_font(r1, size=10, italic=True, bold=True, color=GREY)
    r2 = p.add_run(text)
    set_run_font(r2, size=10, italic=True, color=GREY)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True, color=GREY)
    return p

def add_figure_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'[ {text} ]')
    set_run_font(run, size=10, italic=True, color=GREY)
    return p

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def make_table(doc, headers, rows, caption_text):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, '003DA5')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'EAF4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_run_font(run, size=10)
    add_caption(doc, caption_text)

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'« {text} »')
    set_run_font(run, size=11, italic=True, bold=True, color=AT_BLUE)

def page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('REPUBLIQUE ALGERIENNE DEMOCRATIQUE ET POPULAIRE')
set_run_font(r, size=12, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('MINISTERE DE LA FORMATION ET DE L\'ENSEIGNEMENT PROFESSIONNELS')
set_run_font(r, size=11, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Institut National Spécialisé de la Formation Professionnelle en Audiovisuel')
set_run_font(r, size=11, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run('Echahid Ahmed Mehdi — Ouled Fayet')
set_run_font(r, size=11, italic=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Mémoire De Fin De Formation Pour L\'obtention Du Diplôme De')
set_run_font(r, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Technicien Supérieur En Informatique')
set_run_font(r, size=13, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Option : DEVELOPPEMENT WEB ET MOBILE')
set_run_font(r, size=12, bold=True, color=AT_GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Thème :')
set_run_font(r, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Conception Et Réalisation D\'une Plateforme Web')
set_run_font(r, size=14, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('De Gestion Automatisée Des Missions Et Déplacements')
set_run_font(r, size=14, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Professionnels Au Profit D\'Algérie Télécom')
set_run_font(r, size=14, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run('« AT Réservations »')
set_run_font(r, size=16, bold=True, color=AT_GREEN)

# Info bloc
for label, value in [
    ('Organisme d\'accueil : ', 'Algérie Télécom — Division des Systèmes d\'Information'),
    ('Réalisé par : ',          '[À COMPLÉTER — Nom et Prénom]'),
    ('Suivi par : ',            '[À COMPLÉTER — Encadreur pédagogique, INSFP AV]'),
    ('Encadreur professionnel : ', '[À COMPLÉTER — Nom, Poste, DSI / Algérie Télécom]'),
    ('Promotion : ',            'Avril 2026'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_run_font(r1, size=11, bold=True)
    r2 = p.add_run(value)
    set_run_font(r2, size=11)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  REMERCIEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'Remerciements')
add_paragraph(doc, 'Nous tenons à exprimer notre gratitude envers toutes les personnes qui ont contribué, de près ou de loin, à la réalisation de ce projet. Leur soutien et leurs encouragements ont été précieux tout au long de cette aventure.')
add_paragraph(doc, 'Un merci particulier à nos familles pour leur soutien indéfectible et leurs encouragements constants. Leur confiance en nos capacités a été une source de motivation inestimable.')
add_paragraph(doc, 'Nous remercions notre encadreur pédagogique [À COMPLÉTER] pour ses précieux conseils, sa disponibilité et son suivi tout au long de ce projet, ainsi que notre encadreur professionnel au sein d\'Algérie Télécom [À COMPLÉTER] pour nous avoir accueillis au sein de la Division des Systèmes d\'Information.')
add_paragraph(doc, 'Enfin, nous souhaitons remercier toutes les personnes qui, par leurs contributions directes ou indirectes, ont rendu ce travail possible.')
add_paragraph(doc, '[À COMPLÉTER — Prénom et Nom]', align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)
page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  DEDICACES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'Dédicaces')
add_paragraph(doc, '[À COMPLÉTER PAR L\'ÉTUDIANT]', italic=True, color=GREY)
page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  ABREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'Liste des Abréviations')

abreviations = [
    ('ACTEL',    'Agence Commerciale de Télécommunication'),
    ('API',      'Application Programming Interface'),
    ('AT',       'Algérie Télécom'),
    ('ARPT',     'Autorité de Régulation de la Poste et des Télécommunications'),
    ('CMP',      'Centre de Maintenance et Production'),
    ('CNPE',     'Conseil National aux Participations de l\'État'),
    ('CSS',      'Cascading Style Sheets'),
    ('DOT',      'Direction Opérationnelle des Télécommunications'),
    ('DSI',      'Division des Systèmes d\'Information'),
    ('HTML',     'HyperText Markup Language'),
    ('HTTP',     'HyperText Transfer Protocol'),
    ('INSFP AV', 'Institut National Spécialisé de la Formation Professionnelle en Audiovisuel'),
    ('JSON',     'JavaScript Object Notation'),
    ('MVC',      'Modèle-Vue-Contrôleur'),
    ('MySQL',    'My Structured Query Language'),
    ('PDG',      'Président Directeur Général'),
    ('PHP',      'Hypertext Preprocessor'),
    ('RBAC',     'Role-Based Access Control'),
    ('REST',     'Representational State Transfer'),
    ('SGBD',     'Système de Gestion de Bases de Données'),
    ('SPA',      'Société Par Actions'),
    ('SQL',      'Structured Query Language'),
    ('TS',       'Technicien Supérieur'),
    ('UML',      'Unified Modeling Language'),
    ('URL',      'Uniform Resource Locator'),
]
make_table(doc, ['Sigle', 'Signification'], abreviations, '')
page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'INTRODUCTION')

add_paragraph(doc, "Depuis toujours, les organisations cherchent à améliorer leurs méthodes de gestion et de suivi des activités afin de garantir un fonctionnement efficace et fiable. Au fil du temps, les outils utilisés pour gérer les informations et coordonner les opérations ont considérablement évolué, passant des registres papier et des méthodes manuelles à des systèmes informatiques modernes capables de traiter et de stocker de grandes quantités de données.")
add_paragraph(doc, "Avec l'essor des technologies de l'information et de la communication, les entreprises disposent aujourd'hui de solutions numériques qui facilitent la gestion de leurs activités quotidiennes. Les applications web jouent un rôle particulièrement important dans cette transformation digitale, en permettant un accès rapide et sécurisé aux informations, tout en offrant la possibilité d'automatiser plusieurs processus de gestion. Ces outils contribuent ainsi à améliorer la productivité, la traçabilité des opérations et la prise de décision au sein des organisations.")
add_paragraph(doc, "C'est dans cette perspective que s'inscrit notre projet au sein d'Algérie Télécom, qui souhaite moderniser la gestion de ses missions et déplacements professionnels à travers la mise en place d'une plateforme web centralisée. L'objectif est de disposer d'une solution numérique permettant de remplacer intégralement le circuit manuel, d'améliorer la traçabilité des demandes et d'offrir une meilleure visibilité sur les activités de déplacement de l'ensemble du personnel.")
add_paragraph(doc, "Afin de répondre à ce besoin, notre projet consiste à concevoir et réaliser une application web dédiée à la gestion des missions et déplacements professionnels au sein d'Algérie Télécom, intitulée « AT Réservations ». Cette application permettra aux agents de soumettre leurs demandes en ligne, aux validateurs de les approuver selon le circuit hiérarchique officiel, et aux administrateurs de superviser l'ensemble du processus à travers un tableau de bord analytique.")
add_paragraph(doc, "La réalisation de ce projet nous amène à nous poser la problématique suivante :")
add_blockquote(doc, "Comment concevoir et réaliser une application web au profit d'Algérie Télécom, en atteignant les objectifs assignés et en respectant les contraintes posées ?")
add_paragraph(doc, "Pour répondre à cette problématique, nous avons structuré notre mémoire en quatre chapitres. Le premier chapitre, intitulé « Étude préalable », présente l'organisme d'accueil Algérie Télécom ainsi que le contexte général du projet. Il décrit également la présentation du sujet, la problématique, les objectifs, les contraintes et le public visé.")
add_paragraph(doc, "Le deuxième chapitre, intitulé « Analyse et spécification des besoins », est consacré à l'identification des besoins fonctionnels et non fonctionnels du système, ainsi qu'aux diagrammes UML décrivant le fonctionnement de l'application.")
add_paragraph(doc, "Le troisième chapitre, intitulé « Conception de l'application », présente la conception globale du système, notamment la charte graphique, les maquettes de l'interface utilisateur et le modèle de données.")
add_paragraph(doc, "Le quatrième chapitre, intitulé « Réalisation », décrit les différentes étapes du développement, les technologies utilisées et les tests effectués.")
add_paragraph(doc, "Enfin, nous terminerons ce mémoire par une conclusion générale présentant le bilan du projet ainsi que les perspectives d'amélioration.")
page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPITRE I
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('CHAPITRE I :')
set_run_font(r, size=18, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Étude Préalable')
set_run_font(r, size=16, bold=True, color=AT_GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run('Présentation de l\'organisme d\'accueil, présentation du sujet,\nproblématique, objectifs, contraintes et public visé.')
set_run_font(r, size=11, italic=True, color=GREY)

page_break(doc)

add_heading1(doc, 'CHAPITRE I : Étude Préalable')

# I.1
add_heading2(doc, 'I.1. Présentation de l\'organisme d\'accueil')
add_heading3(doc, 'I.1.1. Présentation d\'Algérie Télécom')
add_paragraph(doc, "Algérie Télécom est une société par actions (SPA) à capitaux publics, leader sur le marché algérien des télécommunications, qui connaît une forte croissance. Elle offre une gamme complète de services de voix et de données aux clients résidentiels et professionnels, et opère sur le marché des réseaux et services de communications électroniques.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run("Sa création a été instituée par la loi n° 2000/03 du 5 août 2000")
set_run_font(r1, size=11)
r2 = p.add_run(' ¹')
set_run_font(r2, size=9, color=GREY)
r3 = p.add_run(", fixant les règles générales relatives à la poste et aux télécommunications. Cette loi a été suivie des résolutions du Conseil National aux Participations de l'État (CNPE) du 1er mars 2001, portant création d'une Entreprise Publique Économique dénommée « Algérie Télécom ».")
set_run_font(r3, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run("Elle est régie par cette loi qui lui confère le statut d'une entreprise publique économique sous la forme juridique d'une société par actions (SPA), au capital social de 115 000 000 000 DA (cent quinze milliards de dinars)")
set_run_font(r1, size=11)
r2 = p.add_run(' ²')
set_run_font(r2, size=9, color=GREY)
r3 = p.add_run(", inscrite au registre du commerce le 11 mai 2002 sous le numéro 02B 0018083.")
set_run_font(r3, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run("Entrée officiellement en activité à partir du 1er janvier 2003")
set_run_font(r1, size=11)
r2 = p.add_run(' ³')
set_run_font(r2, size=9, color=GREY)
r3 = p.add_run(", elle s'est engagée dans le monde des Technologies de l'Information et de la Communication avec trois objectifs fondateurs : la rentabilité, l'efficacité et la qualité de service.")
set_run_font(r3, size=11)

add_heading3(doc, 'I.1.2. Fiche technique d\'Algérie Télécom')
make_table(doc,
    ['Désignation', 'Information'],
    [
        ('Nom de l\'entreprise',   'Algérie Télécom'),
        ('Date de création',        '2003 (entrée en activité officielle)'),
        ('Statut juridique',        'SPA — Société par Actions à capitaux publics'),
        ('Capital social',          '115 000 000 000 DA (cent quinze milliards de dinars)'),
        ('N° Registre commerce',    '02B 0018083'),
        ('Siège social',            'Route Nationale N°05, Cinq Maisons, Mohammadia — Alger'),
        ('Secteur d\'activité',     'Réseaux et services de communications électroniques'),
        ('Présence nationale',      '60 DOT — 146 ACTEL sur l\'ensemble du territoire'),
        ('Site web',                'www.algerietelecom.dz'),
    ],
    'Tableau N°1 : Fiche technique d\'Algérie Télécom'
)

add_heading3(doc, 'I.1.3. Missions et objectifs d\'Algérie Télécom')
add_heading4(doc, 'I.1.3.1. Activité principale')
add_paragraph(doc, "Fournir des services de télécommunication permettant le transport et l'échange de la voix, de messages écrits, de données numériques et d'informations audiovisuelles. Développer, exploiter et gérer les réseaux publics et privés de télécommunications ; établir, exploiter et gérer les interconnexions avec tous les opérateurs des réseaux.")

add_heading4(doc, 'I.1.3.2. Objectifs stratégiques')
add_paragraph(doc, "Algérie Télécom est engagée dans le monde des technologies de l'information et de la communication avec les objectifs suivants :")
add_bullet(doc, "Accroître l'offre de services téléphoniques et faciliter l'accès aux services des télécommunications au plus grand nombre d'usagers, en particulier en zones rurales.")
add_bullet(doc, "Accroître la qualité des services offerts et la gamme des prestations rendues, et rendre plus compétitifs les services de télécommunications.")
add_bullet(doc, "Développer un réseau national de télécommunication fiable et connecté aux autoroutes de l'information.")
add_bullet(doc, "Acquérir de nouvelles parts de marché et devenir un opérateur multimédia.")
add_bullet(doc, "Mettre en place de nouvelles procédures en matière de ressources humaines et de systèmes d'information.")

add_heading3(doc, 'I.1.4. Organisation générale d\'Algérie Télécom')
add_paragraph(doc, "Algérie Télécom est organisée en Divisions, Directions Centrales, et 60 Directions Opérationnelles des Télécommunications (DOT) distribuées sur tout le territoire national, dont trois à Alger. Chaque DOT comprend plusieurs ACTEL (Agences Commerciales de Télécommunication) et CMP (Centres de Maintenance et Production).")
add_figure_placeholder(doc, 'Insérer ici Figure N°1 : Organigramme général d\'Algérie Télécom')
add_caption(doc, 'Figure N°1 : Organigramme général d\'Algérie Télécom')

add_heading3(doc, 'I.1.5. Présentation de la structure d\'accueil : la Division des Systèmes d\'Information (DSI)')
add_paragraph(doc, "Notre stage a été effectué au sein de la Division des Systèmes d'Information (DSI) d'Algérie Télécom. La DSI est une division de services informatiques qui a pour mission de fournir à l'entreprise des systèmes d'information de pointe, couvrant l'ensemble de ses activités. Elle joue un rôle transversal et stratégique en étant l'interlocuteur informatique de toutes les autres divisions et directions.")

add_heading4(doc, 'I.1.5.1. Missions de la DSI')
add_paragraph(doc, "La Division des Systèmes d'Information a pour missions principales :")
add_bullet(doc, "Faire évoluer et maintenir l'infrastructure informatique interne de l'entreprise.")
add_bullet(doc, "Veiller à la pérennité des applications de gestion et à leur intégration dans le système d'information global.")
add_bullet(doc, "Assurer le support technique aux utilisateurs des systèmes d'information sur l'ensemble du territoire national.")
add_bullet(doc, "Gérer et maintenir le tissu informationnel de l'entreprise : archivage, bases de données, portails intranet et documents techniques.")
add_bullet(doc, "Proposer des solutions et services innovants dans le domaine des systèmes d'information pour les clients internes.")
add_bullet(doc, "Constituer et animer un pôle d'expertise dans les systèmes d'information.")
add_figure_placeholder(doc, 'Insérer ici Figure N°2 : Organigramme de la Division des Systèmes d\'Information (DSI)')
add_caption(doc, 'Figure N°2 : Organigramme de la Division des Systèmes d\'Information (DSI)')

# I.2
add_heading2(doc, 'I.2. Présentation du sujet')
add_paragraph(doc, "La plateforme « AT Réservations » est un système de gestion des missions et déplacements professionnels destiné au personnel d'Algérie Télécom. Elle consiste en une solution informatique web intégrée, ayant pour objectif d'assurer une gestion centralisée, sécurisée et efficace de l'ensemble du processus de demande, de validation et de suivi des missions professionnelles.")
add_paragraph(doc, "L'application web constitue le cœur du système et permet aux agents autorisés de créer et soumettre leurs demandes de mission, aux validateurs hiérarchiques de les approuver ou rejeter, et aux administrateurs de superviser l'ensemble des opérations. Elle offre également des fonctionnalités de messagerie interne, de génération de rapports détaillés et de contrôle des accès selon les rôles des utilisateurs, garantissant ainsi la fiabilité et la sécurité des données.")
add_paragraph(doc, "Grâce à cette solution centralisée, Algérie Télécom disposera d'un système moderne permettant d'améliorer la traçabilité des demandes, d'assurer la conformité du circuit de validation hiérarchique et d'optimiser la gestion administrative des missions, tout en renforçant la communication entre les différents niveaux de l'entreprise.")

# I.3
add_heading2(doc, 'I.3. Problématique')
add_paragraph(doc, "Face aux difficultés de gestion des missions professionnelles auxquelles Algérie Télécom est confrontée, telles que l'utilisation de formulaires papier, le manque de centralisation des données, l'absence de visibilité en temps réel sur l'état des demandes, l'absence de mécanismes de traçabilité fiables et la lenteur des circuits de validation manuels, il devient nécessaire de moderniser et d'optimiser le système de gestion actuel.")
add_paragraph(doc, "Ces contraintes engendrent des lenteurs dans le traitement des demandes, des risques élevés de perte de documents, et une difficulté dans le suivi administratif des missions, affectant ainsi l'efficacité globale du fonctionnement de l'entreprise.")
add_paragraph(doc, "D'où notre problématique :")
add_blockquote(doc, "Comment concevoir et réaliser une application web au profit d'Algérie Télécom, en atteignant les objectifs assignés et en respectant les contraintes posées ?")

# I.4
add_heading2(doc, 'I.4. Les objectifs')
add_heading3(doc, 'I.4.1. Objectifs généraux')
add_paragraph(doc, "Nous devons concevoir et développer une application web pour la gestion des missions et déplacements professionnels d'Algérie Télécom, adaptée aux besoins de l'entreprise, tout en répondant aux objectifs suivants :")
add_bullet(doc, "Mettre en œuvre un système de gestion centralisée des demandes de missions et déplacements professionnels.")
add_bullet(doc, "Permettre la gestion automatisée du circuit de validation hiérarchique.")
add_bullet(doc, "Mettre en place un système d'authentification sécurisé basé sur les rôles (administrateur, validateur, demandeur, utilisateur).")
add_bullet(doc, "Assurer le suivi en temps réel de l'état des demandes et des missions en cours.")
add_bullet(doc, "Mettre en œuvre un système de notifications automatiques à chaque étape du processus.")
add_bullet(doc, "Permettre la génération de rapports et statistiques pour les décideurs.")

add_heading3(doc, 'I.4.2. Objectifs par espace utilisateur')

def add_space_section(doc, title, items, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True, color=AT_BLUE)
    for item in items:
        add_bullet(doc, item)
    add_note(doc, 'Note : ', note)

add_space_section(doc, 'Espace Administrateur :',
    [
        "Créer, modifier, activer et désactiver les comptes utilisateurs de tous les agents.",
        "Attribuer et modifier les rôles de chaque utilisateur (demandeur, validateur, administrateur, utilisateur consultatif).",
        "Consulter l'ensemble des demandes de missions, quel que soit leur statut.",
        "Superviser le tableau de bord global avec les statistiques de toutes les directions.",
        "Exporter des rapports complets au format PDF et Excel.",
        "Gérer les paramètres de l'application (structure hiérarchique, organigramme).",
        "Consulter l'historique complet de toutes les validations et actions effectuées.",
        "Envoyer des messages internes à tous les utilisateurs.",
    ],
    "Ce que l'administrateur ne fait pas : il ne crée pas de demandes de mission au nom des agents (ce processus appartient au demandeur)."
)

add_space_section(doc, 'Espace Validateur :',
    [
        "Consulter la liste des demandes de missions en attente de sa validation.",
        "Approuver une demande de mission et faire avancer le circuit hiérarchique.",
        "Rejeter une demande avec un motif obligatoire communiqué au demandeur.",
        "Demander une modification pour renvoyer la demande en correction au demandeur.",
        "Consulter l'historique de toutes ses décisions de validation.",
        "Consulter le tableau de bord de son périmètre (missions de sa direction).",
        "Exporter des rapports concernant les missions de son périmètre.",
        "Envoyer et recevoir des messages internes liés aux demandes en cours.",
    ],
    "Ce que le validateur ne peut pas faire : il ne peut pas créer de demandes, ni modifier les comptes utilisateurs, ni accéder aux données d'autres directions."
)

add_space_section(doc, 'Espace Demandeur :',
    [
        "Créer une nouvelle demande de mission via un formulaire guidé en quatre étapes (informations générales, réservations, documents, récapitulatif).",
        "Enregistrer une demande en brouillon pour la compléter ultérieurement.",
        "Soumettre sa demande au circuit de validation hiérarchique.",
        "Joindre les pièces justificatives nécessaires (ordre de mission, autorisation, etc.).",
        "Consulter l'état d'avancement de toutes ses demandes en temps réel.",
        "Modifier ou corriger une demande renvoyée par le validateur.",
        "Consulter l'historique complet de ses demandes et de leurs validations.",
        "Recevoir des notifications automatiques à chaque changement de statut de sa demande.",
        "Envoyer et recevoir des messages internes liés à ses demandes.",
    ],
    "Ce que le demandeur ne peut pas faire : il ne peut pas valider des demandes, ni accéder aux demandes d'autres agents, ni gérer les comptes utilisateurs."
)

add_space_section(doc, 'Espace Utilisateur consultatif :',
    [
        "Consulter les missions en cours et leurs statuts (sans possibilité de modification).",
        "Consulter le tableau de bord avec les statistiques globales.",
        "Consulter l'organigramme interactif d'Algérie Télécom.",
        "Recevoir des notifications générales concernant l'activité de la plateforme.",
    ],
    "Ce que l'utilisateur consultatif ne peut pas faire : il ne peut ni créer de demandes, ni valider, ni gérer des comptes, ni exporter des données."
)

# I.5
add_heading2(doc, 'I.5. Les contraintes')
add_heading3(doc, 'I.5.1. Contraintes techniques')
add_bullet(doc, "L'application doit être entièrement accessible via un navigateur web standard, sans nécessiter l'installation d'un logiciel client spécifique sur les postes de travail.")
add_bullet(doc, "L'application doit être compatible avec les navigateurs suivants : Mozilla Firefox, Google Chrome, Microsoft Edge.")
add_bullet(doc, "Le déploiement est prévu sur l'infrastructure serveur interne d'Algérie Télécom (Intranet).")
add_bullet(doc, "La base de données doit être hébergée sur un serveur MySQL accessible localement.")

add_heading3(doc, 'I.5.2. Contraintes fonctionnelles')
add_bullet(doc, "Le circuit de validation doit impérativement respecter la structure hiérarchique officielle d'Algérie Télécom.")
add_bullet(doc, "Aucun module de paiement en ligne n'est requis (exclu du cahier des charges).")
add_bullet(doc, "Les notifications par e-mail sont intégrées à l'application ; les notifications SMS sont simulées via des logs applicatifs pour la soutenance.")
add_bullet(doc, "L'application doit fonctionner en mode entièrement opérationnel lors de la soutenance (aucun mode démonstration fictif).")

add_heading3(doc, 'I.5.3. Contraintes de sécurité')
add_bullet(doc, "Mise en place d'un système d'authentification sécurisé et de chiffrement des données pour protéger les informations sensibles.")
add_bullet(doc, "Gestion des rôles et permissions (RBAC — Role-Based Access Control) garantissant qu'aucun utilisateur n'accède à des données dépassant ses attributions.")
add_bullet(doc, "Les mots de passe doivent être stockés sous forme hachée (bcrypt) dans la base de données.")
add_bullet(doc, "Un mécanisme de protection contre les attaques par force brute est appliqué sur la page de connexion (blocage après 5 tentatives échouées).")

add_heading3(doc, 'I.5.4. Contraintes de conception')
add_paragraph(doc, "L'interface utilisateur doit respecter la charte graphique officielle d'Algérie Télécom. Les couleurs institutionnelles à appliquer sont le vert (#00A650) et le bleu (#003DA5). L'ergonomie de l'application doit garantir une navigation intuitive, compatible avec les trois navigateurs ciblés.")

# I.6
add_heading2(doc, 'I.6. Le public visé')
add_paragraph(doc, "L'application « AT Réservations » est destinée à un usage exclusivement interne au sein d'Algérie Télécom. Elle s'adresse aux quatre catégories d'utilisateurs suivantes :")
add_bullet(doc, "Les agents et employés d'Algérie Télécom (profil Demandeur) : toute personne habilitée à initier une demande de mission ou de déplacement professionnel, quelle que soit sa direction d'appartenance.")
add_bullet(doc, "Les responsables hiérarchiques (profil Validateur) : les supérieurs hiérarchiques désignés comme points de validation dans le circuit d'approbation, à différents niveaux de la structure.")
add_bullet(doc, "Les agents de la Division des Systèmes d'Information (profil Administrateur) : chargés de la gestion des comptes utilisateurs, des paramètres applicatifs et de la supervision globale du système.")
add_bullet(doc, "Les cadres et responsables consultatifs (profil Utilisateur) : toute personne autorisée à consulter les données et statistiques sans disposer de droits de modification.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  CHAPITRE II
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('CHAPITRE II :')
set_run_font(r, size=18, bold=True, color=AT_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Analyse et Spécification des Besoins')
set_run_font(r, size=16, bold=True, color=AT_GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run('Concepts théoriques, diagrammes de cas d\'utilisation et diagrammes de séquence')
set_run_font(r, size=11, italic=True, color=GREY)

page_break(doc)

add_heading1(doc, 'CHAPITRE II : Analyse et Spécification des Besoins')

add_heading2(doc, 'Introduction')
add_paragraph(doc, "Ce chapitre présente l'analyse fonctionnelle complète de la plateforme « AT Réservations ». Nous commençons par rappeler quelques concepts théoriques fondamentaux sur UML, puis nous présentons les diagrammes de cas d'utilisation pour chaque acteur du système, et enfin les diagrammes de séquence décrivant les scénarios principaux.")

# II.1
add_heading2(doc, 'II.1. Concepts théoriques')
add_heading3(doc, 'II.1.1. Définition d\'une application web')
add_paragraph(doc, "Une application web est un logiciel qui s'exécute dans un navigateur web. Contrairement aux applications de bureau, elle ne nécessite aucune installation sur le poste client et est accessible depuis n'importe quel appareil connecté. Elle repose sur une architecture client-serveur : le frontend communique avec le backend via des requêtes HTTP.")

add_heading3(doc, 'II.1.2. Le langage UML')
add_paragraph(doc, "Le Langage de Modélisation Unifié (UML) est un langage graphique standardisé utilisé pour visualiser, spécifier et documenter les systèmes logiciels. Il propose différents types de diagrammes adaptés à chaque aspect du système.")

add_heading3(doc, 'II.1.3. Relations dans les diagrammes de cas d\'utilisation')
add_paragraph(doc, "Les diagrammes de cas d'utilisation utilisent deux relations importantes :")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Cm(0.5)
r1 = p.add_run('La relation <<include>> ')
set_run_font(r1, size=11, bold=True, color=AT_BLUE)
r2 = p.add_run("(trait pointillé avec flèche) : un cas d'utilisation inclut obligatoirement un autre. Exemple : « Créer une mission » inclut « S'authentifier » — il est impossible de créer une mission sans être préalablement connecté.")
set_run_font(r2, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Cm(0.5)
r1 = p.add_run('La relation <<extend>> ')
set_run_font(r1, size=11, bold=True, color=AT_BLUE)
r2 = p.add_run("(trait pointillé avec flèche) : indique qu'un cas d'utilisation étend optionnellement un autre, sous certaines conditions seulement. Exemple : « Rejeter avec motif » étend « Traiter une demande » — le rejet n'est déclenché que si le validateur refuse la demande. De même, « Demander une modification » étend « Traiter une demande » — uniquement si la demande doit être renvoyée pour correction.")
set_run_font(r2, size=11)

# II.2
add_heading2(doc, 'II.2. Identification des acteurs')
add_paragraph(doc, "L'application AT Réservations repose sur quatre acteurs principaux :")
make_table(doc,
    ['Acteur', 'Rôle', 'Accès principal'],
    [
        ('Administrateur', 'Super-utilisateur',        'Gestion complète : utilisateurs, prestataires, budgets, audit'),
        ('Validateur',     'Responsable d\'approbation', 'Traitement des demandes : approbation, rejet, demande de modification'),
        ('Demandeur',      'Initiateur des demandes',   'Création, soumission et suivi de ses demandes de mission'),
        ('Utilisateur',    'Lecteur du système',        'Consultation des missions, statistiques et organigramme — aucun droit de création ou de modification'),
    ],
    'Tableau N°2 : Acteurs du système AT Réservations'
)

# II.3
add_heading2(doc, 'II.3. Diagrammes de cas d\'utilisation')
add_heading3(doc, 'II.3.1. Diagramme global du système')
add_paragraph(doc, "Le diagramme suivant représente l'ensemble des fonctionnalités du système et les interactions des quatre acteurs. Tous les cas incluent (<<include>>) « S'authentifier » — aucune fonctionnalité n'est accessible sans connexion préalable.")
add_figure_placeholder(doc, 'Insérer ici Figure N°3 : Diagramme de cas d\'utilisation global — Système AT Réservations')
add_caption(doc, 'Figure N°3 : Diagramme de cas d\'utilisation global — Système AT Réservations')

add_heading3(doc, 'II.3.2. Diagramme de cas d\'utilisation — Administrateur')
add_paragraph(doc, "L'administrateur dispose des accès les plus étendus au sein de la plateforme. Il gère les comptes utilisateurs (activation, désactivation, changement de rôle), les prestataires de services (compagnies aériennes, hôtels, agences de voyage, restauration, transport) ainsi que les budgets par direction et par type de dépense. Il peut consulter les journaux d'audit qui enregistrent toutes les actions critiques du système. Il accède également à l'ensemble des demandes de missions, indépendamment de leur statut ou de leur direction d'origine, et peut générer des exports au format Excel ou PDF.")
add_figure_placeholder(doc, 'Insérer ici Figure N°4 : Diagramme de cas d\'utilisation — Administrateur')
add_caption(doc, 'Figure N°4 : Diagramme de cas d\'utilisation — Administrateur')

add_heading3(doc, 'II.3.3. Diagramme de cas d\'utilisation — Demandeur et Utilisateur')
add_paragraph(doc, "Le demandeur initie des demandes de mission via un formulaire guidé en quatre étapes : (1) Informations générales (titre, destination, dates, type de mission, priorité), (2) Réservations associées (billets d'avion, hébergements, restaurations), (3) Documents justificatifs à joindre, (4) Récapitulatif et soumission. Il peut enregistrer sa demande en brouillon, la modifier, la dupliquer ou l'annuler. Il reçoit des notifications à chaque changement de statut et peut échanger des messages internes avec les autres acteurs du système.")
add_paragraph(doc, "L'utilisateur consultatif dispose d'un accès en lecture seule : il consulte les missions en cours, les statistiques du tableau de bord et l'organigramme interactif d'Algérie Télécom, sans possibilité de créer ni de modifier des données.")
add_figure_placeholder(doc, 'Insérer ici Figure N°5 : Diagramme de cas d\'utilisation — Demandeur et Utilisateur')
add_caption(doc, 'Figure N°5 : Diagramme de cas d\'utilisation — Demandeur et Utilisateur')

add_heading3(doc, 'II.3.4. Diagramme de cas d\'utilisation — Validateur')
add_paragraph(doc, "Le validateur traite exclusivement les demandes de missions qui lui sont assignées dans le circuit de validation. Face à une demande en attente, il dispose de trois actions : approuver la demande, la rejeter avec un motif obligatoire communiqué au demandeur, ou demander une modification pour renvoyer la demande en correction. Les actions de rejet et de demande de modification sont des extensions (<<extend>>) du cas « Traiter une demande ». Le validateur consulte également l'historique de ses décisions et peut exporter les rapports de son périmètre.")
add_figure_placeholder(doc, 'Insérer ici Figure N°6 : Diagramme de cas d\'utilisation — Validateur')
add_caption(doc, 'Figure N°6 : Diagramme de cas d\'utilisation — Validateur')

# II.4
add_heading2(doc, 'II.4. Diagrammes de séquence')
add_heading3(doc, 'II.4.1. Diagramme de séquence — Authentification')
add_paragraph(doc, "Ce diagramme décrit le scénario d'authentification via Laravel Sanctum. L'utilisateur saisit ses identifiants sur le formulaire (Login.jsx). Le service API envoie une requête POST /auth/login vers l'AuthController, qui vérifie les identifiants et applique une protection contre les attaques par force brute (blocage après 5 tentatives échouées). En cas de succès, un token Sanctum est généré et stocké côté client dans le localStorage sous la clé at_token. Le profil de l'utilisateur est ensuite récupéré via GET /auth/me, et la redirection vers le tableau de bord est effectuée selon le rôle détecté.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run('Participants : ')
set_run_font(r1, size=11, bold=True)
r2 = p.add_run("Utilisateur, Login.jsx (Frontend), api.js (Service Axios), AuthController (Backend), Laravel Sanctum (Token), localStorage.")
set_run_font(r2, size=11, italic=True)

add_figure_placeholder(doc, 'Insérer ici Figure N°7 : Diagramme de séquence — Authentification (Laravel Sanctum)')
add_caption(doc, 'Figure N°7 : Diagramme de séquence — Authentification (Laravel Sanctum)')

add_heading3(doc, 'II.4.2. Diagramme de séquence — Création et validation d\'une mission')
add_paragraph(doc, "Ce diagramme représente le scénario complet de création et de validation d'une mission. Le demandeur remplit le formulaire guidé NewMissionWizard.jsx (4 étapes). La mission est créée en base via POST /missions avec le statut « brouillon ». Les réservations (billets d'avion, hébergements, restaurations) sont ajoutées via POST /missions/{id}/reservations. À la soumission (POST /missions/{id}/submit), le statut passe à « soumis », une entrée est créée dans la table circuits_validation et une notification est envoyée au validateur désigné. Le validateur consulte sa liste de demandes en attente, approuve via POST /validations/{id}/approuver, ce qui met à jour le statut de la mission à « approuvé » et envoie une notification ainsi qu'un e-mail au demandeur.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run('Participants : ')
set_run_font(r1, size=11, bold=True)
r2 = p.add_run("Demandeur, NewMissionWizard.jsx (Frontend), API Laravel (Backend), Base de données MySQL (tables : ordres_de_mission, circuits_validation, reservations), Validateur.")
set_run_font(r2, size=11, italic=True)

add_figure_placeholder(doc, 'Insérer ici Figure N°8 : Diagramme de séquence — Créer et valider une demande de mission')
add_caption(doc, 'Figure N°8 : Diagramme de séquence — Créer et valider une demande de mission')

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  WEBOGRAPHIE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'Webographie')
refs = [
    ('¹', 'Loi n° 2000/03 du 5 août 2000 — Journal Officiel de la République Algérienne'),
    ('²', 'Site officiel Algérie Télécom — www.algerietelecom.dz'),
    ('³', 'Groupe Télécom Algérie — fr.wikipedia.org/wiki/Groupe_Télécom_Algérie'),
]
for num, ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f'{num}  ')
    set_run_font(r1, size=11, bold=True, color=AT_BLUE)
    r2 = p.add_run(ref)
    set_run_font(r2, size=11)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/user/AT-r-servation/Memoir_AT_Reservations_CORRIGE.docx'
doc.save(output_path)
print(f'Fichier créé : {output_path}')
